import traceback

from flask import Flask, request, render_template, send_file, jsonify, send_from_directory
import os
import uuid
from pathlib import Path
import logging
import boto3
from io import BytesIO
import logging
import shutil
import openai
import httpx
import time
import tempfile
import subprocess
import concurrent.futures
from datetime import datetime
import sqlalchemy
from sqlalchemy import create_engine, text
from sqlalchemy.exc import SQLAlchemyError
from sqlalchemy.engine.url import make_url
import platform
import sys

# Configure logging first - increase level to see more details
logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Initialise Flask app
application = Flask(__name__)
app = application  # AWS ElasticBeanstalk looks for 'application' while Flask CLI looks for 'app'

# Configure app before importing other modules
application.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev-key-for-local-only')

def create_openai_client():
    """
    Create an OpenAI client with explicit HTTP/1.1 configuration
    to resolve AppRunner connectivity issues.
    """
    # Create a custom HTTPX client with HTTP/1.1 explicitly set
    http_client = httpx.Client(
        http1=True,  # Force HTTP/1.1
        http2=False,  # Disable HTTP/2
        verify=True,  # SSL verification
        timeout=30.0  # Generous timeout
    )

    # Initialize OpenAI client with the custom HTTP client
    client = openai.OpenAI(
        api_key=os.environ.get("OPENAI_API_KEY"),
        http_client=http_client
    )

    return client

# Global client initialization
client = create_openai_client()

# Get DATABASE_URL from environment, with a SQLite fallback
database_url = os.environ.get('DATABASE_URL')
logger.info(f"Database URL from environment: {database_url}")

# Determine which database to use
if not database_url:
    fallback_db = 'sqlite:///fallback.db'
    logger.warning(f"Using fallback database: {fallback_db}")
    application.config['SQLALCHEMY_DATABASE_URI'] = fallback_db
else:
    application.config['SQLALCHEMY_DATABASE_URI'] = database_url

# Now that the database URI is set, we can try to download from S3 if it's SQLite
if application.config['SQLALCHEMY_DATABASE_URI'].startswith('sqlite'):
    db_file = application.config['SQLALCHEMY_DATABASE_URI'].replace('sqlite:///', '')
    s3_bucket = os.environ.get('S3_BUCKET', 'your-backup-bucket-name')

    try:
        # Download DB from S3 if it exists
        s3 = boto3.client('s3')
        s3.download_file(s3_bucket, 'db_backup/' + os.path.basename(db_file), db_file)
        logger.info(f"Downloaded database from S3: {db_file}")
    except Exception as e:
        logger.warning(f"Could not download database from S3: {str(e)}")

# Define backup function to be used later
def backup_db_to_s3():
    if application.config['SQLALCHEMY_DATABASE_URI'].startswith('sqlite'):
        db_file = application.config['SQLALCHEMY_DATABASE_URI'].replace('sqlite:///', '')
        s3_bucket = os.environ.get('S3_BUCKET', 'your-backup-bucket-name')

        try:
            # Upload DB to S3
            s3 = boto3.client('s3')
            s3.upload_file(db_file, s3_bucket, 'db_backup/' + os.path.basename(db_file))
            logger.info(f"Backed up database to S3: {db_file}")
        except Exception as e:
            logger.warning(f"Could not back up database to S3: {str(e)}")

application.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
application.config['TEMPLATES_AUTO_RELOAD'] = True

# Database configuration needs to be engine-specific
if application.config['SQLALCHEMY_DATABASE_URI'].startswith('sqlite'):
    # SQLite-specific options
    application.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
        'connect_args': {'check_same_thread': False}
    }
else:
    # PostgreSQL-specific options
    application.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
        'pool_recycle': 280,
        'pool_timeout': 10,
        'connect_args': {
            'connect_timeout': 5  # PostgreSQL connection timeout in seconds
        }
    }

# Log the database we're connecting to
logger.info(f"Configured database: {application.config['SQLALCHEMY_DATABASE_URI']}")

S3_AVAILABLE = False

try:
    import boto3
    S3_AVAILABLE = True
except ImportError:
    logger.warning("boto3 not available, S3 integration will be disabled")


def backup_db_to_s3():
    if not S3_AVAILABLE:
        logger.warning("S3 not available, skipping backup")
        return False

    if application.config['SQLALCHEMY_DATABASE_URI'].startswith('sqlite'):
        db_file = application.config['SQLALCHEMY_DATABASE_URI'].replace('sqlite:///', '')
        s3_bucket = os.environ.get('S3_BUCKET')

        if not s3_bucket:
            logger.warning("S3_BUCKET not set, skipping backup")
            return False

        try:
            # Upload DB to S3
            s3 = boto3.client('s3')
            s3.upload_file(db_file, s3_bucket, 'db_backup/' + os.path.basename(db_file))
            logger.info(f"Backed up database to S3: {db_file}")
            return True
        except Exception as e:
            logger.warning(f"Could not back up database to S3: {str(e)}")
            return False
    return False

# Make backup function available globally
application.backup_db_to_s3 = backup_db_to_s3

# Import db and login_manager from extensions
from extensions import db, login_manager

# Initialize extensions with the application
db.init_app(application)
login_manager.init_app(application)
login_manager.login_view = 'auth.login'

# Create folders for file storage
UPLOAD_FOLDER = Path('/tmp/temp_uploads')
TRANSCRIPT_FOLDER = Path('/tmp/temp_transcripts')
RESULT_FOLDER = Path('/tmp/results')

try:
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    logging.info(f"Upload folder exists: {os.path.exists(UPLOAD_FOLDER)}")
    logging.info(f"Upload folder permissions: {os.stat(UPLOAD_FOLDER).st_mode}")
except Exception as e:
    logging.error(f"Error checking upload folder: {str(e)}")

for folder in [UPLOAD_FOLDER, TRANSCRIPT_FOLDER, RESULT_FOLDER]:
    folder.mkdir(exist_ok=True, parents=True)


# Create a diagnostic route BEFORE importing models
@application.route('/api/diagnostics')
def diagnostics():
    db_status = "unknown"
    try:
        with db.engine.connect() as conn:
            conn.execute("SELECT 1")
            db_status = "connected"
    except Exception as e:
        db_status = f"error: {str(e)}"

    return jsonify({
        'status': 'ok',
        'message': 'API is running',
        'database': {
            'status': db_status,
            'uri': application.config['SQLALCHEMY_DATABASE_URI'].split('@')[-1].split('/')[0]
            if '@' in application.config['SQLALCHEMY_DATABASE_URI'] else 'local'
        },
        'environment': {
            'FLASK_ENV': os.environ.get('FLASK_ENV', 'production')
        }
    })


# Import models and auth AFTER initializing the database
try:
    from models import User

    logger.info("Models imported successfully")
except Exception as e:
    logger.error(f"Error importing models: {str(e)}")


    # Create a minimal User model for the app to run
    class User(db.Model):
        id = db.Column(db.Integer, primary_key=True)
        username = db.Column(db.String(80), unique=True, nullable=False)
        email = db.Column(db.String(120), unique=True, nullable=False)
        is_active = db.Column(db.Boolean, default=True)
        total_transcription_minutes = db.Column(db.Float, default=0.0)
        current_month_transcription_minutes = db.Column(db.Float, default=0.0)
        last_usage_reset = db.Column(db.DateTime, default=datetime.utcnow)

        def get_id(self):
            return str(self.id)

        @property
        def is_authenticated(self):
            return True

        @property
        def is_anonymous(self):
            return False

# Import auth blueprint
try:
    from auth import auth as auth_blueprint

    application.register_blueprint(auth_blueprint)
    logger.info("Auth blueprint registered successfully")
except Exception as e:
    logger.error(f"Error registering auth blueprint: {str(e)}")
    # Create a minimal auth blueprint for the app to run
    from flask import Blueprint

    auth_blueprint = Blueprint('auth', __name__)
    application.register_blueprint(auth_blueprint)


# Configure user loader
@login_manager.user_loader
def load_user(user_id):
    try:
        return User.query.get(int(user_id))
    except:
        return None


# Import the rest of your dependencies
import speech_recognition as sr
from pydub import AudioSegment
import pandas as pd
import cv2
import numpy as np
import re
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash
import psycopg2
import threading
from concurrent.futures import ThreadPoolExecutor
import multiprocessing
import difflib
from docx import Document
from flask_login import login_required, current_user

# Create database tables ONCE within an application context with improved error handling
try:
    with application.app_context():
        # Test the database connection first
        try:
            logger.info("Testing database connection...")
            with db.engine.connect() as conn:
                conn.execute(text("SELECT 1"))
            logger.info("Database connection successful!")

            # Check if tables already exist before creating them
            logger.info("Checking if database tables exist...")
            try:
                # Try to query the User table to see if it exists
                with db.engine.connect() as conn:
                    result = conn.execute(text("SELECT COUNT(*) FROM \"user\""))
                    user_count = result.scalar()
                logger.info(f"Found existing user table with {user_count} users")
                logger.info("Database tables already exist, skipping creation")
            except Exception as table_check_error:
                # Tables don't exist, so create them
                logger.info("Tables don't exist yet, creating database tables...")
                try:
                    db.create_all()
                    logger.info("Database tables created successfully")
                except Exception as create_error:
                    logger.error(f"Error creating tables: {str(create_error)}")
                    # If we get a "relation already exists" error, that's actually fine
                    if "already exists" in str(create_error).lower():
                        logger.info("Tables already exist (caught during creation), continuing...")
                    else:
                        raise create_error

        except Exception as e:
            logger.error(f"Error with database: {str(e)}")
            logger.warning("Application will continue with limited functionality")
except Exception as e:
    logger.error(f"Database initialization error: {str(e)}")
    logger.warning("Application will run in fallback mode")

def convert_to_wav(audio_path):
    """
    Convert audio file to WAV format if it's not already WAV.
    Returns path to WAV file.
    """
    print(f"Converting {audio_path} to WAV format...")
    audio_path = Path(audio_path)
    if audio_path.suffix.lower() == '.wav':
        return str(audio_path)

    # Create a temporary WAV file
    wav_path = audio_path.with_suffix('.wav')

    # Convert audio to WAV based on file extension
    file_ext = audio_path.suffix.lower()

    try:
        if file_ext == '.mp3':
            audio = AudioSegment.from_mp3(str(audio_path))
        elif file_ext == '.flac':
            audio = AudioSegment.from_file(str(audio_path), format="flac")
        elif file_ext == '.m4a':
            audio = AudioSegment.from_file(str(audio_path), format="m4a")
        elif file_ext == '.aac':
            audio = AudioSegment.from_file(str(audio_path), format="aac")
        else:
            # Try to automatically detect format
            audio = AudioSegment.from_file(str(audio_path))

        audio.export(wav_path, format="wav")
        print(f"Successfully converted to {wav_path}")
        return str(wav_path)
    except Exception as e:
        print(f"Error converting {audio_path}: {str(e)}")
        # If conversion fails, try to return the original path
        # Speech recognition might still work for some formats
        return str(audio_path)


def optimise_audio_file(file_path, output_dir=None, quality='speech', min_silence_len=700):
    """
    Optimise audio file for transcription by applying multiple compression techniques.

    Args:
        file_path: Path to the original audio file
        output_dir: Directory to save the optimised file (defaults to same directory)
        quality: Optimization profile ('speech' or 'high')
        min_silence_len: Minimum silence length in ms to detect and compress

    Returns:
        Tuple of (path to optimised file, compression stats dictionary)
    """
    start_time = time.time()
    file_path = Path(file_path)

    if output_dir is None:
        output_dir = file_path.parent
    else:
        output_dir = Path(output_dir)
        output_dir.mkdir(exist_ok=True, parents=True)

    # Generate output filename
    output_filename = f"optimised_{file_path.stem}.mp3"
    output_path = output_dir / output_filename

    # Log start of optimization
    logging.info(f"Starting audio optimization for: {file_path}")
    original_size = file_path.stat().st_size
    logging.info(f"Original file size: {original_size / 1024 / 1024:.2f} MB")

    try:
        # Load audio file
        audio = AudioSegment.from_file(str(file_path))
        original_duration = len(audio) / 1000  # in seconds

        # Step 1: Convert to mono if stereo
        was_stereo = audio.channels > 1
        if was_stereo:
            audio = audio.set_channels(1)
            logging.info("Converted audio to mono")

        # Step 2: Downsample to 16kHz for speech (22kHz for high quality)
        target_rate = 16000 if quality == 'speech' else 22050
        if audio.frame_rate != target_rate:
            audio = audio.set_frame_rate(target_rate)
            logging.info(f"Downsampled audio to {target_rate}Hz")

        # Step 3: Detect and compress silence
        if min_silence_len > 0:
            # Detect silence
            silence_threshold = -40  # dB, adjust as needed
            chunks = detect_silence_chunks(audio, silence_threshold, min_silence_len)

            if chunks:
                # Compress silence by keeping just enough to maintain natural pauses
                compressed_silence_len = 300  # ms for each silence
                audio = compress_silence(audio, chunks, compressed_silence_len)
                logging.info(f"Compressed {len(chunks)} silence segments")

        # Step 4: Export with optimised encoding
        # For speech, use very aggressive compression
        if quality == 'speech':
            bitrate = "24k"  # Very low bitrate suitable for speech
        else:
            bitrate = "64k"  # Higher quality for general audio

        # Export to WAV with specified parameters
        audio.export(
            output_path,
            format="mp3",
            parameters=[
                "-ar", str(target_rate),
                "-ac", "1",
                "-b:a", bitrate
            ]
        )

        # Calculate compression stats
        end_time = time.time()
        compressed_size = output_path.stat().st_size
        size_reduction = 1 - (compressed_size / original_size)
        compressed_duration = len(audio) / 1000  # in seconds
        duration_reduction = 1 - (compressed_duration / original_duration)

        # Log results
        logging.info(f"Optimised file saved to: {output_path}")
        logging.info(f"Optimised file size: {compressed_size / 1024 / 1024:.2f} MB")
        logging.info(f"Size reduction: {size_reduction:.2%}")
        logging.info(f"Duration reduction: {duration_reduction:.2%}")
        logging.info(f"Optimization completed in {end_time - start_time:.2f} seconds")

        # Return stats dictionary
        stats = {
            "original_path": str(file_path),
            "optimised_path": str(output_path),
            "original_size": original_size,
            "optimised_size": compressed_size,
            "size_reduction_percent": size_reduction * 100,
            "original_duration": original_duration,
            "optimised_duration": compressed_duration,
            "duration_reduction_percent": duration_reduction * 100,
            "processing_time": end_time - start_time,
            "was_stereo": was_stereo,
            "target_sample_rate": target_rate,
            "silence_segments_compressed": len(chunks) if min_silence_len > 0 else 0
        }

        return output_path, stats

    except Exception as e:
        logging.error(f"Error optimizing audio: {str(e)}")
        return None, {"error": str(e)}


def detect_silence_chunks(audio, silence_threshold=-40, min_silence_len=700):
    """
    Detect silent segments in audio.

    Args:
        audio: AudioSegment to analyze
        silence_threshold: dB threshold below which is considered silence
        min_silence_len: minimum silence length in ms

    Returns:
        List of (start_ms, end_ms) tuples of silence chunks
    """
    # Get audio data as numpy array
    samples = np.array(audio.get_array_of_samples())
    sample_rate = audio.frame_rate

    # Convert to float between -1 and 1
    if audio.sample_width == 2:  # 16-bit
        max_value = 32768.0
    elif audio.sample_width == 1:  # 8-bit
        max_value = 128.0
    else:  # 24 or 32-bit
        max_value = 2 ** (8 * audio.sample_width - 1)

    samples = samples / max_value

    # Calculate dB values (using small epsilon to avoid log(0))
    epsilon = 1e-10
    db_values = 20 * np.log10(np.maximum(np.abs(samples), epsilon))

    # Find segments below threshold
    is_silence = db_values < silence_threshold

    # Convert to milliseconds positions
    silence_chunks = []
    in_silence = False
    silence_start = 0

    # Factor to convert sample position to ms
    ms_per_sample = 1000.0 / sample_rate

    for i, silent in enumerate(is_silence):
        position_ms = i * ms_per_sample

        if silent and not in_silence:
            # Start of silence
            in_silence = True
            silence_start = position_ms
        elif not silent and in_silence:
            # End of silence
            silence_end = position_ms
            silence_duration = silence_end - silence_start

            # Only keep if long enough
            if silence_duration >= min_silence_len:
                silence_chunks.append((int(silence_start), int(silence_end)))

            in_silence = False

    # Check if file ends in silence
    if in_silence:
        silence_end = len(samples) * ms_per_sample
        silence_duration = silence_end - silence_start

        if silence_duration >= min_silence_len:
            silence_chunks.append((int(silence_start), int(silence_end)))

    return silence_chunks


def compress_silence(audio, silence_chunks, target_silence_len=300):
    """
    Compress silence segments in audio to a target length.

    Args:
        audio: AudioSegment to modify
        silence_chunks: List of (start_ms, end_ms) tuples of silence
        target_silence_len: Target length in ms for each silence segment

    Returns:
        New AudioSegment with compressed silence
    """
    if not silence_chunks:
        return audio

    # Sort chunks by start time
    silence_chunks.sort(key=lambda x: x[0])

    # Initialize new audio segments list
    segments = []
    last_end = 0

    for start, end in silence_chunks:
        # Add audio before silence
        if start > last_end:
            segments.append(audio[last_end:start])

        # Add compressed silence - keep the middle portion
        silence_duration = end - start
        if silence_duration > target_silence_len:
            # Calculate start and end points for the portion to keep
            keep_start = start + (silence_duration - target_silence_len) // 2
            keep_end = keep_start + target_silence_len
            segments.append(audio[keep_start:keep_end])
        else:
            # If silence is already shorter than target, keep it all
            segments.append(audio[start:end])

        last_end = end

    # Add remaining audio after last silence
    if last_end < len(audio):
        segments.append(audio[last_end:])

    # Concatenate all segments
    if segments:
        return sum(segments)
    return audio


def process_uploaded_audio(file_path, original_filename):
    """
    Process and optimize an uploaded audio file.

    Args:
        file_path: Path to the uploaded audio file
        original_filename: Original filename

    Returns:
        Tuple of (optimized_path, error_message)
    """
    # Create a separate folder for optimized files
    optimised_folder = Path('/tmp/optimized_audio')
    optimised_folder.mkdir(exist_ok=True, parents=True)

    # Log the process
    logging.info(f"Processing uploaded file: {original_filename}")

    try:
        # Optimize the audio
        optimised_path, stats = optimise_audio_file(
            file_path,
            output_dir=optimised_folder,
            quality='speech',  # Optimize for speech
            min_silence_len=700  # Min silence length to compress (ms)
        )

        if optimised_path is None:
            return None, f"Failed to optimize audio: {stats.get('error', 'Unknown error')}"

        # Log the optimization results
        logging.info(
            f"Audio optimization for {original_filename}: {stats['size_reduction_percent']:.2f}% size reduction")

        # Return the optimized file path
        return optimised_path, None

    except Exception as e:
        logging.error(f"Error processing audio {original_filename}: {str(e)}")
        return None, f"Error processing audio: {str(e)}"


def transcribe_audio_optimized(audio_path):
    """
    Optimized version that preserves structure markers in audio transcription.
    Prioritizes structure over speed.
    """
    try:
        print(f"Processing {audio_path} for transcription...")
        audio_path = Path(audio_path)

        # Verify file exists
        if not audio_path.exists():
            print(f"Error: File does not exist: {audio_path}")
            return ""

        # Convert audio to WAV format with good quality for recognition
        wav_path = audio_path.with_suffix('.wav')
        if audio_path.suffix.lower() != '.wav':
            try:
                audio = AudioSegment.from_file(str(audio_path))
                # Use higher quality settings for better accuracy
                audio = audio.set_frame_rate(16000).set_channels(1)
                audio.export(wav_path, format="wav")
                print(f"Successfully converted to {wav_path}")
            except Exception as e:
                print(f"Error converting audio: {e}")
                return ""
        else:
            wav_path = str(audio_path)

        # Get audio duration
        audio_segment = AudioSegment.from_file(str(wav_path))
        duration_seconds = len(audio_segment) / 1000.0
        print(f"Audio duration: {duration_seconds:.2f} seconds")

        # CRITICAL CHANGE: Try to transcribe the entire audio first
        # This gives the best chance of preserving all markers
        try:
            print("Attempting full audio transcription without chunking...")
            recognizer = sr.Recognizer()

            with sr.AudioFile(str(wav_path)) as source:
                audio_data = recognizer.record(source)

                # Try with a longer timeout for the entire file
                try:
                    import signal

                    def handler(signum, frame):
                        raise TimeoutError("Full transcription timed out")

                    # Set a generous timeout based on audio length
                    timeout_seconds = min(30, max(10, int(duration_seconds * 0.5)))
                    print(f"Setting timeout of {timeout_seconds} seconds for full transcription")

                    signal.signal(signal.SIGALRM, handler)
                    signal.alarm(timeout_seconds)

                    with open(wav_path, "rb") as audio_file:
                        whisper_response = client.audio.transcriptions.create(
                            model="whisper-1",
                            file=audio_file,
                            response_format="text",
                            language="en"
                        )
                        transcript = whisper_response

                    # Clear the alarm
                    signal.alarm(0)

                    print("Full transcription successful!")
                    print(f"Transcript length: {len(transcript)} characters")

                    # If successful, return the full transcript
                    return transcript.lower()

                except TimeoutError:
                    print("Full transcription timed out, falling back to chunking")
                    signal.alarm(0)  # Clear the alarm
                except Exception as e:
                    print(f"Error in full transcription: {e}")
                    signal.alarm(0)  # Clear the alarm
        except Exception as e:
            print(f"Error attempting full transcription: {e}")

        # If we reach here, full transcription failed - try larger, overlapping chunks
        print("Using overlapping chunks for transcription...")

        # Use larger chunks with overlap to avoid splitting markers
        chunk_size_ms = 15000  # 15 seconds
        overlap_ms = 5000  # 5 second overlap

        # Calculate actual chunks with overlap
        chunks = []
        for start_ms in range(0, len(audio_segment), chunk_size_ms - overlap_ms):
            end_ms = min(start_ms + chunk_size_ms, len(audio_segment))
            chunks.append((start_ms, end_ms))

        print(f"Processing audio in {len(chunks)} overlapping chunks")

        # Process each chunk
        all_transcripts = []
        recognizer = sr.Recognizer()

        for i, (start_ms, end_ms) in enumerate(chunks):
            # Extract chunk with overlap
            chunk = audio_segment[start_ms:end_ms]

            # Export chunk to temporary file
            chunk_path = f"/tmp/chunk_{i}.wav"
            chunk.export(chunk_path, format="wav")

            print(f"Processing chunk {i + 1}/{len(chunks)}...")

            try:
                with sr.AudioFile(chunk_path) as source:
                    audio_data = recognizer.record(source)

                    def handler(signum, frame):
                        raise TimeoutError("Chunk processing timed out")

                    # Set timeout proportional to chunk size
                    timeout_sec = min(10, int((end_ms - start_ms) / 1000 * 0.8))
                    signal.signal(signal.SIGALRM, handler)
                    signal.alarm(timeout_sec)

                    try:
                        transcript = recognizer.recognize_google(audio_data, language="en-US")
                        all_transcripts.append(transcript)
                        signal.alarm(0)  # Clear the alarm
                        print(f"Chunk {i + 1}/{len(chunks)} transcribed successfully")
                    except TimeoutError:
                        print(f"Chunk {i + 1}/{len(chunks)} timed out")
                        signal.alarm(0)  # Clear the alarm
                        all_transcripts.append("")
                    except Exception as e:
                        print(f"Error with chunk {i + 1}: {e}")
                        signal.alarm(0)  # Clear the alarm
                        all_transcripts.append("")
            except Exception as e:
                print(f"Error processing chunk {i + 1}: {e}")
            finally:
                # Clean up chunk file
                try:
                    os.remove(chunk_path)
                except:
                    pass

        # Combine all transcripts into a single string
        raw_transcript = " ".join(all_transcripts).lower()

        # CRITICAL STEP: Process the raw transcript to ensure markers are preserved
        # This is where we clean up and ensure our special markers are correctly identified
        processed_transcript = fix_transcript_markers(raw_transcript)

        return processed_transcript

    except Exception as e:
        print(f"Transcription error: {e}")
        import traceback
        traceback.print_exc()
        return ""
    finally:
        # Clean up temporary WAV file
        if 'wav_path' in locals() and wav_path != str(audio_path) and Path(wav_path).exists():
            try:
                Path(wav_path).unlink(missing_ok=True)
                print("Cleaned up temporary WAV file")
            except:
                pass


def fix_transcript_markers(transcript):
    """
    Carefully process the transcript to ensure structure markers are correctly identified.
    This is critical for proper CSV structure.
    """
    # Lower case for consistent processing
    transcript = transcript.lower()

    # Fix common speech recognition errors in markers
    marker_corrections = {
        # "new room" variations
        r'\bnew\s*rum\b': 'new room',
        r'\bnew\s*rome\b': 'new room',
        r'\bnew\s*rooms\b': 'new room',
        r'\bnew\s*rom\b': 'new room',
        r'\bnu\s*room\b': 'new room',

        # "new attribute" variations
        r'\bnew\s*attribute\b': 'new attribute',
        r'\bnew\s*attributes\b': 'new attribute',
        r'\bnew\s*attributed\b': 'new attribute',
        r'\bnew\s*tribute\b': 'new attribute',

        # "new feature" variations
        r'\bnew\s*feature\b': 'new feature',
        r'\bnew\s*features\b': 'new feature',
        r'\bnew\s*featured\b': 'new feature',
        r'\bnew\s*creature\b': 'new feature',

        # "comment" variations
        r'\bcomment\b': 'comment',
        r'\bcomments\b': 'comment',
        r'\bcomet\b': 'comment',
        r'\bcoming\b': 'comment'
    }

    # Apply corrections to fix misheard markers
    for pattern, replacement in marker_corrections.items():
        transcript = re.sub(pattern, replacement, transcript)

    # Ensure spaces around markers for better detection
    transcript = re.sub(r'(new room|new attribute|new feature|comment)', r' \1 ', transcript)

    # Clean up excess whitespace
    transcript = re.sub(r'\s+', ' ', transcript).strip()

    # Split into words and rebuild transcript carefully
    words = transcript.split()
    processed_words = []

    # Track special markers
    i = 0
    while i < len(words):
        if i < len(words) - 1 and words[i] == 'new' and words[i + 1] in ['room', 'attribute', 'feature']:
            # Found a marker like "new room", "new attribute", "new feature"
            processed_words.append(f"\n{words[i]} {words[i + 1]}")
            i += 2
        elif words[i] == 'comment':
            # Found a "comment" marker
            processed_words.append(f"\ncomment")
            i += 1
        else:
            # Regular word
            processed_words.append(words[i])
            i += 1

    # Join back into a clean transcript with line breaks at markers
    processed_transcript = " ".join(processed_words)

    # Clean up any excess spaces
    processed_transcript = re.sub(r'\s+', ' ', processed_transcript)

    # Ensure markers are at line starts
    processed_transcript = re.sub(r' (new room|new attribute|new feature|comment)', r'\n\1', processed_transcript)

    # Remove any empty lines
    lines = [line.strip() for line in processed_transcript.split('\n') if line.strip()]
    processed_transcript = '\n'.join(lines)

    print(f"Processed transcript structure: {len(lines)} lines with markers")

    return processed_transcript


def process_audio_in_chunks(wav_path, audio_segment=None):
    """
    Process longer audio files in chunks to avoid timeouts.
    """
    import concurrent.futures  # Explicit import here as well

    if audio_segment is None:
        audio_segment = AudioSegment.from_file(str(wav_path))

    # Define chunk size (10 seconds)
    chunk_size_ms = 10000  # 10 seconds

    # Calculate number of chunks
    duration_ms = len(audio_segment)
    num_chunks = (duration_ms // chunk_size_ms) + (1 if duration_ms % chunk_size_ms > 0 else 0)

    print(f"Processing audio in {num_chunks} chunks of 10 seconds each")

    # Initialize recognizer
    recognizer = sr.Recognizer()
    recognizer.energy_threshold = 300
    recognizer.dynamic_energy_threshold = True
    recognizer.pause_threshold = 0.8

    # Process each chunk
    all_transcripts = []

    for i in range(num_chunks):
        start_ms = i * chunk_size_ms
        end_ms = min((i + 1) * chunk_size_ms, duration_ms)

        # Extract chunk
        chunk = audio_segment[start_ms:end_ms]

        # Export chunk to temporary file
        chunk_path = f"/tmp/chunk_{i}.wav"
        chunk.export(chunk_path, format="wav")

        print(f"Processing chunk {i + 1}/{num_chunks}...")

        try:
            with sr.AudioFile(chunk_path) as source:
                audio_data = recognizer.record(source)

                try:
                    # Use a short timeout for each chunk to prevent worker timeouts
                    with concurrent.futures.ThreadPoolExecutor() as executor:
                        future = executor.submit(recognizer.recognize_google, audio_data, language="en-US")
                        transcript = future.result(timeout=5)  # 5 second timeout per chunk
                    print(f"Chunk {i + 1}/{num_chunks} transcribed successfully")
                    all_transcripts.append(transcript)
                except concurrent.futures.TimeoutError:
                    print(f"Chunk {i + 1}/{num_chunks} timed out")
                    all_transcripts.append("")  # Add blank for timed out chunk
                except Exception as e:
                    print(f"Error processing chunk {i + 1}: {e}")
                    all_transcripts.append("")  # Add blank for failed chunk
        except Exception as e:
            print(f"Error processing chunk {i + 1}: {e}")
        finally:
            # Clean up temporary chunk file
            try:
                os.remove(chunk_path)
            except:
                pass

    # Join all transcripts
    full_transcript = " ".join(all_transcripts)
    print(f"Completed chunked transcription, total length: {len(full_transcript)} characters")

    return full_transcript.lower()


def track_transcription_usage(file_path, user_id=None):
    """
    Calculate audio duration and update the user's transcription usage statistics.
    Args:
        file_path: Path to the audio file
        user_id: Optional user ID (if not provided, will try to get from current_user)
    Returns:
        True if successful, False otherwise
    """
    try:
        # Calculate audio duration
        audio = AudioSegment.from_file(str(file_path))
        duration_minutes = len(audio) / 60000  # Convert milliseconds to minutes
        print(f"Audio duration: {duration_minutes:.2f} minutes")

        # Get user ID if not provided
        if user_id is None:
            try:
                from flask_login import current_user
                from flask import has_request_context

                if has_request_context() and current_user and not current_user.is_anonymous:
                    user_id = current_user.id
                else:
                    print("No user context available for usage tracking")
                    return False
            except Exception as e:
                print(f"Error getting current user: {str(e)}")
                return False

        # Now get the user and update stats
        user = User.query.get(user_id)
        if not user:
            print(f"User with ID {user_id} not found")
            return False

        # Check if user model has the required attributes
        required_attrs = ['total_transcription_minutes', 'current_month_transcription_minutes', 'last_usage_reset']
        for attr in required_attrs:
            if not hasattr(user, attr):
                print(f"User model missing attribute: {attr}")
                return False

        # Update total usage
        user.total_transcription_minutes += duration_minutes

        # Check if we need to reset the monthly counter
        current_time = datetime.utcnow()
        if not user.last_usage_reset or user.last_usage_reset.month != current_time.month or user.last_usage_reset.year != current_time.year:
            # New month - reset counter
            user.current_month_transcription_minutes = duration_minutes
            user.last_usage_reset = current_time
        else:
            # Same month - add to counter
            user.current_month_transcription_minutes += duration_minutes

        # Commit changes safely
        try:
            db.session.commit()
            print(f"Updated usage for user {user_id}: added {duration_minutes:.2f} minutes")
            return True
        except Exception as e:
            db.session.rollback()
            print(f"Database error updating usage: {str(e)}")
            return False

    except Exception as e:
        print(f"Error tracking transcription usage: {str(e)}")
        try:
            db.session.rollback()
        except:
            pass
        return False

def format_text(text):
    """
    Apply advanced formatting rules to make transcribed text more professional.
    Uses heuristics to detect sentence boundaries and format accordingly.

    Args:
        text: Raw text string from transcription

    Returns:
        Formatted text with proper capitalization and punctuation
    """
    if not text:
        return text

    # Preprocessing - clean up common transcription artifacts
    text = text.strip()

    # Replace multiple spaces with a single space
    text = ' '.join(text.split())

    # Sentence boundary detection
    # Add periods where they're likely missing between sentences

    # Common indicators of sentence boundaries in speech
    boundary_indicators = [
        ' and then ', ' afterwards ', ' after that ', ' next ', ' following that ',
        ' subsequently ', ' consequently ', ' as a result ', ' therefore ',
        ' thus ', ' hence ', ' so ', ' finally ', ' lastly ', ' in conclusion ',
        ' to sum up ', ' in summary ', ' additionally ', ' moreover ', ' furthermore ',
        ' in addition ', ' also ', ' besides ', ' however ', ' nevertheless ',
        ' nonetheless ', ' still ', ' yet ', ' conversely ', ' on the other hand ',
        ' on the contrary ', ' in contrast ', ' instead ', ' alternatively ',
        ' otherwise ', ' rather ', ' whereas ', ' while ', ' though ', ' although '
    ]

    # Add periods before these indicators if there's no punctuation
    for indicator in boundary_indicators:
        text = text.replace(indicator, f'. {indicator.strip()}')

    # Replace common conjunctions at the start of a sentence that might indicate a boundary
    conjunctions = [' but ', ' or ', ' nor ', ' for ', ' so ']
    for conj in conjunctions:
        text = text.replace(conj, f'. {conj.strip()} ')

    # Split text into sentences based on existing punctuation and the indicators we added
    # This regex will split on ., !, ? followed by a space or end of string
    import re
    sentences = re.split(r'([.!?])\s+', text)

    # Rejoin sentences with proper spacing and formatting
    formatted_text = ''
    i = 0
    while i < len(sentences):
        if i < len(sentences) - 1 and sentences[i + 1] in '.!?':
            # This is a sentence + its ending punctuation
            sentence = sentences[i]
            punct = sentences[i + 1]

            # Capitalize first letter of sentence if it's not already
            if sentence and sentence[0].islower():
                sentence = sentence[0].upper() + sentence[1:]

            formatted_text += sentence + punct + ' '
            i += 2
        else:
            # This might be the last sentence or one without punctuation
            sentence = sentences[i]

            # Capitalize first letter of sentence if it's not already
            if sentence and sentence[0].islower():
                sentence = sentence[0].upper() + sentence[1:]

            # Add period if there's no ending punctuation
            if sentence and not sentence.endswith(('.', '!', '?')):
                sentence += '.'

            formatted_text += sentence + ' '
            i += 1

    # Final cleanup and formatting
    formatted_text = formatted_text.strip()

    # Fix common capitalization issues
    formatted_text = re.sub(r'\bi\b', 'I', formatted_text)  # Capitalize standalone 'i'
    formatted_text = re.sub(r'\bi\'m\b', 'I\'m', formatted_text, flags=re.IGNORECASE)
    formatted_text = re.sub(r'\bi\'ll\b', 'I\'ll', formatted_text, flags=re.IGNORECASE)
    formatted_text = re.sub(r'\bi\'ve\b', 'I\'ve', formatted_text, flags=re.IGNORECASE)
    formatted_text = re.sub(r'\bi\'d\b', 'I\'d', formatted_text, flags=re.IGNORECASE)

    # Capitalize proper nouns (common ones in your context)
    proper_nouns = ['monday', 'tuesday', 'wednesday', 'thursday', 'friday', 'saturday', 'sunday',
                    'january', 'february', 'march', 'april', 'may', 'june', 'july', 'august',
                    'september', 'october', 'november', 'december']

    for noun in proper_nouns:
        formatted_text = re.sub(r'\b' + noun + r'\b', noun.capitalize(),
                                formatted_text, flags=re.IGNORECASE)

    # Fix spacing around punctuation
    formatted_text = re.sub(r'\s+([.,;:!?])', r'\1', formatted_text)

    # Ensure consistent spacing after punctuation
    formatted_text = re.sub(r'([.,;:!?])(\S)', r'\1 \2', formatted_text)

    # Remove double periods
    formatted_text = formatted_text.replace('..', '.')

    # Ensure the text ends with proper punctuation
    if not formatted_text[-1] in '.!?':
        formatted_text += '.'

    return formatted_text

def debug_raw_transcript(audio_path):
    """Debug function to show raw transcription before any processing."""
    transcript = transcribe_audio_optimized(audio_path)
    print("\n--- RAW TRANSCRIPT ---")
    print(transcript)
    print("--- END RAW TRANSCRIPT ---\n")
    return transcript

def correct_transcript_with_gpt(transcript):
    """
    Use GPT-3.5-Turbo to correct property inspection transcription errors
    while strictly preserving the format.

    Args:
        transcript: The raw transcription text from basic speech recognition

    Returns:
        Corrected transcript text with domain-specific terms fixed
    """
    try:
        # Skip processing if transcript is empty
        if not transcript:
            return transcript

        print(f"Starting GPT correction of transcript...")
        start_time = time.time()

        # Log the original transcript for debugging
        logging.info(f"Original transcript: {transcript}")

        # Create a system prompt with property inspection domain knowledge
        # with strict instructions about format preservation
        system_prompt = """You are an expert assistant for property inspectors. 
        Your task is to correct property inspection transcription errors while precisely maintaining the exact format.

        EXTREMELY IMPORTANT: You must preserve all format markers EXACTLY as they appear:
        - "new room" must remain exactly as "new room"
        - "new attribute" must remain exactly as "new attribute"
        - "new feature" must remain exactly as "new feature" 
        - "comment" must remain exactly as "comment"

        These markers control how the transcript is processed, and any change to them will break the processing.

        Only correct real estate terminology, proper nouns, and obvious speech recognition errors:
        - "Germany" should be corrected to "generally" when referring to condition
        - "Mark" to "mark" when not referring to a person's name
        - "tennant" to "tenant"
        - "lanlord" to "landlord"

        DO NOT add punctuation, capitalization, or restructure the text in any way.
        DO NOT combine or split sections marked by these key phrases.
        DO NOT add or remove any "new room", "new attribute", "new feature", or "comment" markers.

        Return ONLY the corrected transcript with minimal changes.
        """

        # Send the transcript to GPT-3.5-Turbo for correction
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user",
                 "content": f"Correct ONLY the terminology errors in this property inspection transcript while preserving the exact format:\n\n{transcript}"}
            ],
            temperature=0.1,  # Very low temperature for more deterministic output
            max_tokens=2000
        )

        # Extract the corrected text
        corrected_transcript = response.choices[0].message.content.strip()

        # Log the corrected transcript for debugging
        logging.info(f"Corrected transcript: {corrected_transcript}")

        # Check if key markers are preserved
        key_markers = ["new room", "new attribute", "new feature", "comment"]
        for marker in key_markers:
            original_count = transcript.lower().count(marker)
            corrected_count = corrected_transcript.lower().count(marker)

            if original_count != corrected_count:
                logging.warning(
                    f"Format marker '{marker}' count changed! Original: {original_count}, Corrected: {corrected_count}")
                return transcript  # Return original if format changed

        return corrected_transcript

    except Exception as e:
        logging.error(f"Error in GPT correction: {str(e)}")
        # Fall back to original transcript if API call fails
        return transcript


def correct_csv_with_gpt(csv_path):
    """
    Process a completed transcript CSV with GPT to correct terminology and speech errors.

    Args:
        csv_path: Path to the CSV file containing transcription data

    Returns:
        Path to the corrected CSV file
    """
    try:
        # Read the CSV file
        df = pd.read_csv(csv_path)

        # Check if we have comments to process
        if 'Comment' not in df.columns:
            logging.warning(f"No 'Comment' column found in {csv_path}")
            return csv_path

        # Process each comment with GPT-4
        logging.info(f"Processing {len(df)} comments with GPT-4")
        for i, row in df.iterrows():
            if pd.notna(row['Comment']) and row['Comment'].strip():
                original_comment = row['Comment']
                corrected_comment = correct_transcript_with_gpt(original_comment)

                # Update the dataframe with corrected comment
                df.at[i, 'Comment'] = corrected_comment

                # Log if there was a change
                if original_comment != corrected_comment:
                    logging.info(f"Correction made: '{original_comment}' → '{corrected_comment}'")

        # Create path for corrected file
        original_filename = Path(csv_path).name
        corrected_filename = f"corrected_{original_filename}"
        corrected_path = Path(csv_path).parent / corrected_filename

        # Save the corrected CSV
        df.to_csv(corrected_path, index=False)
        logging.info(f"Corrected CSV saved to {corrected_path}")

        return corrected_path

    except Exception as e:
        logging.error(f"Error processing CSV with GPT: {str(e)}")
        # Return original path if processing fails
        return csv_path


# Alternative approach: Post-process existing CSVs
def post_process_existing_csv(csv_path):
    """Apply GPT correction to an existing CSV file."""
    return correct_csv_with_gpt(csv_path)


def process_transcript(transcript):
    """
    Process transcript text following exact hierarchical rules:
    * New Room → new row with room, first attribute, first feature, first comment
    * New Attribute → new row with attribute, first feature, first comment
    * New Feature → new row with feature, first comment
    * New Comment → new row with just the comment

    Updated to handle continuous text with trigger words embedded as space-separated phrases.
    """
    print("\nProcessing transcript...")
    results = []

    # Clean up the transcript
    transcript = transcript.strip().lower()

    # Split the transcript at trigger words to create segments
    # Use regex to split while keeping the trigger words
    import re

    # Define trigger patterns
    trigger_pattern = r'(new room|new attribute|new feature|comment)'

    # Split the transcript while keeping the delimiters
    segments = re.split(trigger_pattern, transcript)

    # Remove empty segments and strip whitespace
    segments = [seg.strip() for seg in segments if seg.strip()]

    print(f"Split transcript into {len(segments)} segments: {segments}")

    # Process segments in pairs (trigger + content)
    current_room = ""
    current_attribute = ""
    current_feature = ""
    pending_row = {
        "Room": "",
        "Attribute": "",
        "Feature": "",
        "Comment": "",
        "Tenant Responsibility (TR)": ""
    }

    i = 0
    while i < len(segments):
        segment = segments[i].strip()

        # Skip empty segments
        if not segment:
            i += 1
            continue

        # Check if this segment is a trigger word
        if segment in ['new room', 'new attribute', 'new feature', 'comment']:
            trigger = segment

            # Get the content following this trigger (next segment)
            content = ""
            if i + 1 < len(segments):
                content = segments[i + 1].strip()
                i += 2  # Skip both trigger and content
            else:
                i += 1  # Just skip the trigger if no content follows

            print(f"Processing trigger: '{trigger}' with content: '{content}'")

            # Handle each trigger type
            if trigger == "new room":
                # Add any pending row first
                if any(pending_row.values()):
                    results.append(pending_row.copy())
                    print(f"Added pending row: {pending_row}")

                # Start a new row with this room
                current_room = content
                pending_row = {
                    "Room": current_room,
                    "Attribute": "",
                    "Feature": "",
                    "Comment": "",
                    "Tenant Responsibility (TR)": ""
                }
                print(f"Started new room: '{current_room}'")

            elif trigger == "new attribute":
                # If we have a pending room but no attribute yet, add to same row
                if pending_row["Room"] and not pending_row["Attribute"]:
                    current_attribute = content
                    pending_row["Attribute"] = current_attribute
                    print(f"Added attribute to existing row: '{current_attribute}'")
                else:
                    # Otherwise, save pending row and start new row
                    if any(pending_row.values()):
                        results.append(pending_row.copy())
                        print(f"Added pending row: {pending_row}")

                    current_attribute = content
                    pending_row = {
                        "Room": "",  # Empty for hierarchical display
                        "Attribute": current_attribute,
                        "Feature": "",
                        "Comment": "",
                        "Tenant Responsibility (TR)": ""
                    }
                    print(f"Started new attribute row: '{current_attribute}'")

            elif trigger == "new feature":
                # If we have a pending attribute but no feature yet, add to same row
                if (pending_row["Room"] or pending_row["Attribute"]) and not pending_row["Feature"]:
                    current_feature = content
                    pending_row["Feature"] = current_feature
                    print(f"Added feature to existing row: '{current_feature}'")
                else:
                    # Otherwise, save pending row and start new row
                    if any(pending_row.values()):
                        results.append(pending_row.copy())
                        print(f"Added pending row: {pending_row}")

                    current_feature = content
                    pending_row = {
                        "Room": "",  # Empty for hierarchical display
                        "Attribute": "",  # Empty for hierarchical display
                        "Feature": current_feature,
                        "Comment": "",
                        "Tenant Responsibility (TR)": ""
                    }
                    print(f"Started new feature row: '{current_feature}'")

            elif trigger == "comment":
                is_tenant_responsibility = "tenant responsibility" in content.lower()

                # Apply text formatting if function is available
                if callable(format_text):
                    try:
                        content = format_text(content)
                    except Exception as e:
                        print(f"Warning: Error formatting comment: {str(e)}")

                # If we have a pending feature but no comment yet, add to same row
                if (pending_row["Room"] or pending_row["Attribute"] or pending_row["Feature"]) and not pending_row[
                    "Comment"]:
                    pending_row["Comment"] = content
                    pending_row["Tenant Responsibility (TR)"] = "✓" if is_tenant_responsibility else ""
                    print(f"Added comment to existing row: '{content}'")
                else:
                    # Otherwise, save pending row and start new row with just the comment
                    if any(pending_row.values()):
                        results.append(pending_row.copy())
                        print(f"Added pending row: {pending_row}")

                    pending_row = {
                        "Room": "",  # Empty for hierarchical display
                        "Attribute": "",  # Empty for hierarchical display
                        "Feature": "",  # Empty for hierarchical display
                        "Comment": content,
                        "Tenant Responsibility (TR)": "✓" if is_tenant_responsibility else ""
                    }
                    print(f"Started new comment row: '{content}'")
        else:
            # This segment is not a trigger word, skip it
            # (This shouldn't happen with our regex split, but just in case)
            i += 1

    # Add final pending row if any
    if any(pending_row.values()):
        results.append(pending_row.copy())
        print(f"Added final pending row: {pending_row}")

    print(f"Processed transcript into {len(results)} rows")

    # Debug: print each result row
    for idx, row in enumerate(results):
        print(f"Row {idx + 1}: {row}")

    return results


import csv


def process_audio_file(file_path, original_filename):
    """
    Process audio file with enhanced GPT correction.
    Consolidated version of previous duplicated functions.
    """
    # Get original transcription
    transcript = transcribe_audio_optimized(file_path)

    if not transcript:
        return None, "Failed to transcribe audio"

    # Apply GPT correction
    corrected_transcript = correct_transcript_with_gpt(transcript)

    # Process the corrected transcript
    results = process_transcript(corrected_transcript)

    if not results:
        return None, "No features found in the transcript"

    # Create DataFrame
    df = pd.DataFrame(results)

    # Generate output filename
    base_filename = Path(original_filename).stem
    output_filename = f"{base_filename}_transcript.csv"
    output_path = TRANSCRIPT_FOLDER / output_filename

    # Save to CSV with consistent quoting
    df.to_csv(output_path, index=False, quoting=1)  # 1 is equivalent to csv.QUOTE_ALL

    return output_path, None


# Add a basic test route
@application.route('/test')
def test():
    return "Server is working!"


@application.route('/basic')
def basic():
    return """
    <!DOCTYPE html>
    <html>
    <head>
        <title>PhillyScript Test</title>
        <style>
            body { font-family: sans-serif; padding: 20px; }
            h1 { color: #4a6fa5; }
        </style>
    </head>
    <body>
        <h1>PhillyScript - Test Page</h1>
        <p>If you can see this page, your Flask server is working correctly.</p>
    </body>
    </html>
    """

@application.route('/')
def index():
    try:
        # Serve the index.html page as the root
        return render_template('index.html')
    except Exception as e:
        return f"Error: {str(e)}"

# Add a separate route for the transcription page
@application.route('/transcribe')
@login_required
def transcribe():
    try:
        return render_template('transcription_page.html')
    except Exception as e:
        return f"Error: {str(e)}"


@application.route('/upload', methods=['POST'])
def upload_file():
    if 'audioFile' not in request.files:
        return jsonify({'status': 'error', 'message': 'No file part'})

    file = request.files['audioFile']

    if file.filename == '':
        return jsonify({'status': 'error', 'message': 'No selected file'})

    # Check if file has an allowed extension
    allowed_extensions = ['.mp3', '.wav', '.flac', '.m4a', '.aac']
    file_ext = Path(file.filename).suffix.lower()

    if file_ext not in allowed_extensions:
        return jsonify({
            'status': 'error',
            'message': 'Unsupported file type. Please upload an MP3, WAV, FLAC, M4A, or AAC file.'
        })

    # Generate a unique ID for this upload
    process_id = str(uuid.uuid4())

    # Create temp file path
    temp_file_path = UPLOAD_FOLDER / f"{process_id}_{file.filename}"

    # Save the file
    file.save(temp_file_path)

    # Return process ID so client can check progress
    return jsonify({
        'status': 'success',
        'processId': process_id,
        'message': 'File uploaded successfully. Processing started.'
    })


@application.route('/process/<process_id>', methods=['POST'])
def process_file(process_id):
    """Process uploaded audio file with optimized transcription"""
    # Find the uploaded file
    uploaded_files = list(UPLOAD_FOLDER.glob(f"{process_id}_*"))

    if not uploaded_files:
        return jsonify({'status': 'error', 'message': 'No file found for this process ID'})

    file_path = uploaded_files[0]
    original_filename = file_path.name.replace(f"{process_id}_", "")

    try:
        # First optimize the audio file
        optimized_path, error_msg = process_uploaded_audio(str(file_path), original_filename)

        if error_msg:
            return jsonify({'status': 'error', 'message': error_msg})

        # Get the transcript with preserved markers
        print(f"Starting transcription...")
        transcript = transcribe_audio_optimized(str(optimized_path))

        if not transcript:
            return jsonify({'status': 'error', 'message': 'Failed to transcribe audio'})

        # Apply GPT correction to improve transcript quality
        print(f"Applying GPT correction...")
        corrected_transcript = correct_transcript_with_gpt(transcript)

        # Verify the transcript has the required markers
        marker_count = sum(1 for marker in ['new room', 'new attribute', 'new feature', 'comment']
                           if marker in transcript.lower())

        print(f"Transcript obtained, {marker_count} markers detected")

        # Log the raw transcript for debugging
        print("\n--- RAW TRANSCRIPT ---")
        print(transcript)
        print("--- END RAW TRANSCRIPT ---\n")

        # Process the transcript
        results = process_transcript(corrected_transcript)

        if not results:
            return jsonify({'status': 'error', 'message': 'No features found in the transcript'})

        print(f"Processed into {len(results)} CSV rows")

        # Create DataFrame
        df = pd.DataFrame(results)

        # Generate output filename
        base_filename = Path(original_filename).stem
        output_filename = f"{base_filename}_transcript.csv"
        output_path = TRANSCRIPT_FOLDER / output_filename

        # Save to CSV
        df.to_csv(output_path, index=False)

        return jsonify({
            'status': 'success',
            'outputFilename': Path(output_path).name,
            'message': f'Processing complete - created {len(results)} rows'
        })

    except Exception as e:
        print(f"Error processing file: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'status': 'error', 'message': str(e)})
    finally:
        # Clean up uploaded file
        try:
            file_path.unlink(missing_ok=True)
        except:
            pass


@application.route('/correct_csv/<filename>', methods=['POST'])
@login_required
def correct_csv(filename):
    """Route to apply GPT correction to an existing CSV file."""
    try:
        # Find the CSV file
        csv_path = TRANSCRIPT_FOLDER / filename
        if not csv_path.exists():
            return jsonify({
                'status': 'error',
                'message': 'CSV file not found'
            })

        # Apply GPT correction
        corrected_path = post_process_existing_csv(csv_path)

        # Return path to corrected file
        return jsonify({
            'status': 'success',
            'outputFilename': Path(corrected_path).name,
            'message': 'CSV corrected successfully'
        })
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': f'Error correcting CSV: {str(e)}'
        })

@application.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    file_path = TRANSCRIPT_FOLDER / filename
    if not file_path.exists():
        return jsonify({'status': 'error', 'message': 'File not found'})

    # Return the file for download
    return send_file(file_path, as_attachment=True)


@application.route('/diff_check')
@login_required
def diff_check():
    """Render the image difference checker page"""
    return render_template('diff_check.html')


@application.route('/diff_result/<result_id>')
def diff_result(result_id):
    """Render the results page for image comparison"""
    return render_template('diff_result.html')


@application.route('/api/check_files/<result_id>')
def check_files(result_id):
    """Check if result files exist"""
    original_path = RESULT_FOLDER / f"{result_id}_original.jpg"
    diff_path = RESULT_FOLDER / f"{result_id}_diff.jpg"

    return jsonify({
        'result_folder': str(RESULT_FOLDER),
        'original_exists': original_path.exists(),
        'diff_exists': diff_path.exists(),
        'files_in_folder': [f.name for f in RESULT_FOLDER.iterdir() if f.is_file()][:10]  # List first 10 files
    })
@application.route('/results/<filename>')
def serve_result_file(filename):
    """Serve files from the results folder"""
    return send_from_directory(RESULT_FOLDER, filename)

@application.route('/api/diff_result/<result_id>')
def get_diff_result(result_id):
    """API endpoint to get the image comparison results"""
    # Check if result exists
    original_path = RESULT_FOLDER / f"{result_id}_original.jpg"
    diff_path = RESULT_FOLDER / f"{result_id}_diff.jpg"

    if not original_path.exists() or not diff_path.exists():
        return jsonify({
            'status': 'error',
            'message': 'Result not found'
        })

    # Read the metadata file for difference count
    metadata_path = RESULT_FOLDER / f"{result_id}_metadata.txt"
    difference_count = 0
    if metadata_path.exists():
        with open(metadata_path, 'r') as f:
            metadata = f.read().strip()
            try:
                difference_count = int(metadata)
            except:
                difference_count = 0

    # Return paths and metadata with UPDATED URL PATHS
    return jsonify({
        'status': 'success',
        'originalImageUrl': f"/results/{result_id}_original.jpg",  # Updated path
        'diffImageUrl': f"/results/{result_id}_diff.jpg",         # Updated path
        'differenceCount': difference_count
    })


@application.route('/compare_images', methods=['POST'])
def compare_images():
    """Endpoint to handle image comparison uploads and processing"""
    if 'image1' not in request.files or 'image2' not in request.files:
        return jsonify({
            'status': 'error',
            'message': 'Both images are required'
        })

    file1 = request.files['image1']
    file2 = request.files['image2']

    if file1.filename == '' or file2.filename == '':
        return jsonify({
            'status': 'error',
            'message': 'No file selected'
        })

    # Generate a unique ID for this comparison
    result_id = str(uuid.uuid4())

    # Save uploaded files temporarily
    temp_path1 = UPLOAD_FOLDER / f"{result_id}_1{Path(secure_filename(file1.filename)).suffix}"
    temp_path2 = UPLOAD_FOLDER / f"{result_id}_2{Path(secure_filename(file2.filename)).suffix}"

    file1.save(temp_path1)
    file2.save(temp_path2)

    # Process images and find differences
    try:
        difference_count = process_image_difference(
            str(temp_path1),
            str(temp_path2),
            str(RESULT_FOLDER / f"{result_id}_original.jpg"),
            str(RESULT_FOLDER / f"{result_id}_diff.jpg")
        )

        # Save metadata
        with open(RESULT_FOLDER / f"{result_id}_metadata.txt", 'w') as f:
            f.write(str(difference_count))

        # Clean up temporary files
        temp_path1.unlink(missing_ok=True)
        temp_path2.unlink(missing_ok=True)

        return jsonify({
            'status': 'success',
            'resultId': result_id,
            'message': 'Images compared successfully'
        })

    except Exception as e:
        # Clean up temporary files
        temp_path1.unlink(missing_ok=True)
        temp_path2.unlink(missing_ok=True)

        return jsonify({
            'status': 'error',
            'message': f'Error processing images: {str(e)}'
        })


def process_image_difference(image1_path, image2_path, output_original_path, output_diff_path):
    """
    Process two images to find and highlight differences

    Args:
        image1_path: Path to the reference image
        image2_path: Path to the comparison image
        output_original_path: Path where to save the original image
        output_diff_path: Path where to save the diff image with highlighted differences

    Returns:
        The number of differences found
    """
    # Load images
    img1 = cv2.imread(image1_path)
    img2 = cv2.imread(image2_path)

    # Check if images were loaded correctly
    if img1 is None or img2 is None:
        raise ValueError("Failed to load one or both images")

    # Resize images to match if they have different dimensions
    if img1.shape != img2.shape:
        # Resize the second image to match the first
        img2 = cv2.resize(img2, (img1.shape[1], img1.shape[0]))

    # Convert images to grayscale
    gray1 = cv2.cvtColor(img1, cv2.COLOR_BGR2GRAY)
    gray2 = cv2.cvtColor(img2, cv2.COLOR_BGR2GRAY)

    # Compute absolute difference between grayscale images
    diff = cv2.absdiff(gray1, gray2)

    # Apply threshold to highlight differences
    _, thresh = cv2.threshold(diff, 30, 255, cv2.THRESH_BINARY)

    # Find contours of differences
    contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    # Filter contours by size (remove very small differences that might be noise)
    min_area = 50  # Minimum area in pixels
    significant_contours = [cnt for cnt in contours if cv2.contourArea(cnt) > min_area]

    # Create a copy of the second image to draw the differences on
    img_diff = img2.copy()

    # Draw rectangles around detected differences
    for contour in significant_contours:
        x, y, w, h = cv2.boundingRect(contour)
        cv2.rectangle(img_diff, (x, y), (x + w, y + h), (0, 255, 0), 2)

    # Save the images
    cv2.imwrite(output_original_path, img1)
    cv2.imwrite(output_diff_path, img_diff)

    return len(significant_contours)


@application.route('/finalise_report')
@login_required
def finalise_report():
    """Render the finalise report page"""
    return render_template('finalise_report.html')


@application.route('/compare_text', methods=['POST'])
@login_required
def compare_text():
    """Endpoint to handle text comparison uploads and processing"""
    if 'original' not in request.files or 'comparison' not in request.files:
        return jsonify({
            'status': 'error',
            'message': 'Both files are required'
        })

    original_file = request.files['original']
    comparison_file = request.files['comparison']

    if original_file.filename == '' or comparison_file.filename == '':
        return jsonify({
            'status': 'error',
            'message': 'No file selected'
        })

    # Generate a unique ID for this comparison
    result_id = str(uuid.uuid4())

    # Save uploaded files temporarily
    original_path = UPLOAD_FOLDER / f"{result_id}_original{Path(secure_filename(original_file.filename)).suffix}"
    comparison_path = UPLOAD_FOLDER / f"{result_id}_comparison{Path(secure_filename(comparison_file.filename)).suffix}"

    original_file.save(original_path)
    comparison_file.save(comparison_path)

    # Process files and find differences
    try:
        # Extract text from files
        original_text = extract_text_from_file(original_path)
        comparison_text = extract_text_from_file(comparison_path)

        if not original_text or not comparison_text:
            raise ValueError("Could not extract text from one or both files")

        # Compare texts and generate HTML with differences highlighted
        result_html = generate_diff_html(original_text, comparison_text)

        # Save the result
        result_path = RESULT_FOLDER / f"{result_id}_diff.html"
        with open(result_path, 'w', encoding='utf-8') as f:
            f.write(result_html)

        # Clean up temporary files
        original_path.unlink(missing_ok=True)
        comparison_path.unlink(missing_ok=True)

        return jsonify({
            'status': 'success',
            'resultUrl': f"/view_comparison/{result_id}",  # Changed URL to point to a new route
            'message': 'Files compared successfully'
        })

    except Exception as e:
        # Clean up temporary files
        original_path.unlink(missing_ok=True)
        comparison_path.unlink(missing_ok=True)

        return jsonify({
            'status': 'error',
            'message': f'Error processing files: {str(e)}'
        })


@application.route('/view_comparison/<result_id>')
def view_comparison(result_id):
    """Serve the comparison result directly"""
    result_path = RESULT_FOLDER / f"{result_id}_diff.html"

    if not result_path.exists():
        return "Comparison result not found", 404

    with open(result_path, 'r', encoding='utf-8') as f:
        html_content = f.read()

    return html_content


@application.route('/generate_report', methods=['POST'])
@login_required
def generate_report():
    """Endpoint to handle report generation from CSV"""
    use_latest = request.form.get('useLatest') == 'true'
    report_type = request.form.get('reportType', 'full')  # Default to full if not specified

    try:
        csv_path = None

        if use_latest:
            # Find the most recent CSV file in the transcripts folder
            transcript_dir = Path('/tmp/temp_transcripts')
            if not transcript_dir.exists():
                return jsonify({
                    'status': 'error',
                    'message': 'No transcript directory found'
                })

            csv_files = list(transcript_dir.glob('*.csv'))
            if not csv_files:
                return jsonify({
                    'status': 'error',
                    'message': 'No CSV files found in the transcript directory'
                })

            # Sort by modification time, newest first
            csv_files.sort(key=lambda x: x.stat().st_mtime, reverse=True)
            csv_path = csv_files[0]
        else:
            if 'csv' not in request.files:
                return jsonify({
                    'status': 'error',
                    'message': 'No CSV file provided'
                })

            csv_file = request.files['csv']
            if csv_file.filename == '':
                return jsonify({
                    'status': 'error',
                    'message': 'No file selected'
                })

            # Generate a unique ID for this report
            result_id = str(uuid.uuid4())

            # Save uploaded file temporarily
            csv_path = UPLOAD_FOLDER / f"{result_id}_{secure_filename(csv_file.filename)}"
            csv_file.save(csv_path)

        # Generate report from CSV with the specified report type
        # For now, all report types use the same generation function
        # In the future, you can implement different generation logic based on report_type

        # Log the report type being generated
        logger.info(f"Generating report of type: {report_type}")

        result_path = generate_enhanced_docx_report(csv_path, report_type)

        # Clean up temporary file if it was uploaded
        if not use_latest:
            csv_path.unlink(missing_ok=True)

        return jsonify({
            'status': 'success',
            'resultUrl': f"/download_report/{result_path.name}",
            'message': f'Report generated successfully'
        })

    except Exception as e:
        # Clean up temporary file if it exists and was uploaded
        if csv_path and not use_latest and csv_path.exists():
            csv_path.unlink(missing_ok=True)

        logger.error(f"Error generating report: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': f'Error generating report: {str(e)}'
        })


def generate_report_boilerplate(doc, report_type, address, inspection_date, on_behalf_of):
    """
    Generate standardized boilerplate content for different report types.

    Args:
        doc: The Document object to add content to
        report_type: Type of report ('inventory', 'full', or 'checkout')
        address: Property address
        inspection_date: Date of inspection (formatted string)
        on_behalf_of: Preparer name
    """
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Add logo (placeholder)
    # In a real implementation, you would add a company logo here
    logo_paragraph = doc.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_run = logo_paragraph.add_run("Ft² Inventories")
    logo_run.bold = True
    logo_run.font.size = Pt(16)

    # Add address on the right side
    address_paragraph = doc.add_paragraph()
    address_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    address_run = address_paragraph.add_run(address)

    # Add main report titles based on report type
    if report_type == 'inventory':
        # INVENTORY CHECK-IN REPORT FORMAT
        title = doc.add_heading('INVENTORY', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add And
        and_heading = doc.add_heading('And', 0)
        and_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add CHECK IN
        checkin_heading = doc.add_heading('CHECK IN', 0)
        checkin_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add subtitle
        subtitle = doc.add_paragraph('Of the contents and conditions for')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add property address in a box
        address_box_para = doc.add_paragraph()
        address_box_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        address_box = address_box_para.add_run(address)

        # Add date of inspection
        date_para = doc.add_paragraph(f"Date of inspection: {inspection_date}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add prepared by
        prepared_para = doc.add_paragraph(f"PREPARED BY: Ft² Inventories Ltd")
        prepared_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add company address
        company_address = doc.add_paragraph("Birch Tree House, Glympton Road, Wootton, Woodstock, OX20 1EJ")
        company_address.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add company contact
        company_contact = doc.add_paragraph("info@ft2inventories.co.uk  020 8004 3324")
        company_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add on behalf of section
        behalf_para = doc.add_paragraph("On behalf of:")
        behalf_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add name in a table
        behalf_table = doc.add_table(rows=1, cols=2)
        behalf_table.autofit = False
        behalf_table.columns[0].width = Inches(1)
        behalf_table.columns[1].width = Inches(5)
        behalf_cells = behalf_table.rows[0].cells
        behalf_cells[0].text = "Name"
        behalf_cells[1].text = on_behalf_of

        # Add important information heading
        info_heading = doc.add_heading('Important Information', level=1)
        info_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add important information text
        info_text = doc.add_paragraph(
            "The Tenant /Tenant's representative and landlord should sign this document to signify that they have read and understood each page and accept that this Inventory and Check in is a true and accurate representation at the date so specified of the Property at the address stated above. Amendments to this document may be made within 5 days from the date the report is sent out. If no amendments are made within this time the parties will be deemed to accept the document as an accurate representation at the date so specified of the property at the address so stated above without the need for signature. This document should be returned, signed and dated no later than 5 days from the date the report is sent out.")

        # Add table of contents heading
        toc_heading = doc.add_heading('Table of Contents', level=1)

        # Add placeholder table of contents
        toc = doc.add_table(rows=6, cols=2)
        toc.style = 'Table Grid'

        # Add table of contents rows
        toc_rows = [
            ("Contents", "Page number"),
            ("Disclaimer and important information", "[tc1]"),
            ("General description of conditions/ Utility readings", "[tc2]"),
            ("Front Door and Entrance Hall", "[tc3]"),
            ("Bedroom", "[tc4]"),
            ("Bathroom", "[tc5]")
        ]

        for i, (content, page) in enumerate(toc_rows):
            cells = toc.rows[i].cells
            cells[0].text = content
            cells[1].text = page

        # Add a page break before disclaimer section
        doc.add_page_break()

        # Add INVENTORY DISCLAIMERS heading
        disclaimer_heading = doc.add_heading('INVENTORY DISCLAIMERS', level=1)
        disclaimer_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add disclaimer table
        disclaimer_table = doc.add_table(rows=10, cols=2)
        disclaimer_table.style = 'Table Grid'

        # Add disclaimer content (just a few examples)
        disclaimer_rows = [
            ("1.", "Structural", "This Inventory does not constitute a structural survey of the Property."),
            ("2.", "General",
             "This Inventory has been prepared on the accepted principle that all items are free from any obvious damage, fault or soiling except where stated. The term 'good' is noted as a guideline for this. NOTE: Should there be any cleaning, health and safety or other issues that need urgent attention, please call the Agent's office immediately."),
            ("3.", "Description",
             "Where the words 'gold', 'brass', 'oak', 'walnut' etc are used, it is understood that this is a description of the colour and type of the item and not the actual fabric, unless documentary evidence is available."),
            ("4.", "Attendees",
             "The Clerk must be alone in the property for the inventory inspection. The tenant / Landlord or a representative is welcome to briefly attend the inspection should they wish to do so."),
            ("5.", "Fire Safety Equipment",
             "If smoke detectors/carbon monoxide monitors are present and replacement batteries are required between maintenance visits or periodic tenancy checks, it is the Tenant's responsibility to replace and frequently check the working order of the same.")
        ]

        for i, (num, title, desc) in enumerate(disclaimer_rows[:5]):
            cells = disclaimer_table.rows[i].cells
            cells[0].text = title
            cells[1].text = desc

    elif report_type == 'full':
        # FULL CHECK-IN REPORT FORMAT
        title = doc.add_heading('Full Check-In Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add property address
        address_para = doc.add_paragraph(address)
        address_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add date
        date_paragraph = doc.add_paragraph()
        date_run = date_paragraph.add_run(f"Inspection Date: {inspection_date}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add prepared by
        prepared_para = doc.add_paragraph(f"Prepared by: {on_behalf_of}")
        prepared_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add introduction
        doc.add_heading('Introduction', level=1)
        doc.add_paragraph(
            "This full check-in report documents the complete condition of the property at the beginning of the tenancy and identifies any issues that require attention. Items marked with 'TR' indicate tenant responsibility.")

    elif report_type == 'checkout':
        # CHECK-OUT REPORT FORMAT
        title = doc.add_heading('Check-Out Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add property address
        address_para = doc.add_paragraph(address)
        address_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add date
        date_paragraph = doc.add_paragraph()
        date_run = date_paragraph.add_run(f"Inspection Date: {inspection_date}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add prepared by
        prepared_para = doc.add_paragraph(f"Prepared by: {on_behalf_of}")
        prepared_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add introduction
        doc.add_heading('Introduction', level=1)
        doc.add_paragraph(
            "This check-out report documents the condition of the property at the end of the tenancy and identifies any changes from the check-in report. Items marked with 'TR' indicate tenant responsibility.")

    # Add page break after boilerplate
    doc.add_page_break()

    # Add footer with date to every page
    # Note: This would require more complex handling with python-docx
    # For now, we'll just add a placeholder at the bottom of each page
    footer_para = doc.add_paragraph(f"Date of inspection: {inspection_date}")
    footer_para.style = 'Footer'


def generate_report_closing(doc, report_type, page_count):
    """
    Generate standardized closing content for different report types.

    Args:
        doc: The Document object to add content to
        report_type: Type of report ('inventory', 'full', or 'checkout')
        page_count: Total number of pages in the document (for declaration)
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Add a page break before closing sections
    doc.add_page_break()

    # Add closing content based on report type
    if report_type in ['inventory', 'full']:
        # Add tenant and landlord declaration section
        doc.add_heading('TENANT DECLARATION', level=1)
        doc.add_paragraph(
            f"The items listed in all [x-pages] of this inventory/check in have been inspected and found to be in the condition indicated.")
        doc.add_paragraph(
            "I ………………………………………. being of sound mind have fully understood the implications of signing this document and verify that the content is correct and accurate at the time of signing and that the content will be binding if relied upon in a Court of Law.")

        # Add signature fields
        doc.add_paragraph("Name …………………………………………………………")
        doc.add_paragraph("Signed for the Tenant ……………………………………………")

        doc.add_heading('LANDLORD DECLARATION', level=1)
        doc.add_paragraph(
            f"The items listed in all 124 pages of this inventory/check in have been inspected and found to be in the condition indicated.")
        doc.add_paragraph(
            "I ………………………………………. being of sound mind have fully understood the implications of signing this document and verify that the content is correct and accurate at the time of signing and that the content will be binding if relied upon in a Court of Law.")

        # Add signature fields
        doc.add_paragraph("Name ……………………………………………………………")
        doc.add_paragraph("Signed for the Landlord …………………………………………")

    elif report_type == 'checkout':
        # Add summary
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(
            "This check-out report provides a comprehensive overview of the property's condition at the end of the tenancy. Any discrepancies with the check-in report have been noted.")

        # Add signatures for checkout
        doc.add_heading('CHECKOUT CONFIRMATION', level=1)
        doc.add_paragraph(
            "This checkout report has been completed and represents the condition of the property at the end of the tenancy.")

        # Add signature fields
        doc.add_paragraph("Tenant Name: …………………………………………………………")
        doc.add_paragraph("Tenant Signature: ……………………………………………")
        doc.add_paragraph("Date: ……………………………………")

        doc.add_paragraph("Agent/Landlord Name: …………………………………………………………")
        doc.add_paragraph("Agent/Landlord Signature: ……………………………………………")
        doc.add_paragraph("Date: ……………………………………")

@application.route('/download_report/<filename>')
def download_report(filename):
    """Serve the generated report for download"""
    return send_from_directory(RESULT_FOLDER, filename, as_attachment=True)


def extract_text_from_file(file_path):
    """Extract text content from various file formats"""
    file_ext = file_path.suffix.lower()

    try:
        if file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()

        elif file_ext in ['.doc', '.docx']:
            # Use the properly imported Document class
            from docx import Document
            doc = Document(file_path)
            return '\n'.join([para.text for para in doc.paragraphs])

        elif file_ext == '.pdf':
            # Note: This would require PyPDF2 or another PDF library
            # For simplicity, we'll just return an error for now
            raise ValueError("PDF extraction not implemented")

        elif file_ext == '.rtf':
            # RTF parsing is complex, you'd need a library like striprtf
            raise ValueError("RTF extraction not implemented")

        else:
            # Try to read as plain text by default
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()

    except Exception as e:
        raise ValueError(f"Could not extract text from file: {str(e)}")


def generate_diff_html(original_text, comparison_text):
    """
    Generate HTML highlighting differences between two texts
    """
    # Split texts into lines
    original_lines = original_text.splitlines()
    comparison_lines = comparison_text.splitlines()

    # Get differences
    differ = difflib.HtmlDiff(wrapcolumn=80)
    diff_html = differ.make_file(original_lines, comparison_lines,
                                 'Original Text', 'Comparison Text',
                                 context=True, numlines=3)

    # Enhance the HTML with our own styling
    styled_html = f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Text Comparison Results</title>
        <style>
            body {{
                font-family: 'Segoe UI', Arial, sans-serif;
                margin: 20px;
                line-height: 1.5;
            }}
            .diff-header {{
                background-color: #4a6fa5;
                color: white;
                padding: 10px;
                border-radius: 5px;
                margin-bottom: 20px;
                display: flex;
                justify-content: space-between;
                align-items: center;
            }}
            .diff-header h1 {{
                margin: 0;
                font-size: 24px;
            }}
            table.diff {{
                width: 100%;
                border-collapse: collapse;
                font-size: 14px;
                font-family: monospace;
            }}
            .diff td {{
                padding: 5px;
                border: 1px solid #ddd;
                vertical-align: top;
            }}
            .diff span.diff_add {{
                background-color: #e6ffe6;
                color: green;
                font-weight: bold;
            }}
            .diff span.diff_sub {{
                background-color: #ffe6e6;
                color: red;
                text-decoration: line-through;
            }}
            .diff span.diff_chg {{
                background-color: #fff5cc;
                color: #996600;
                font-weight: bold;
            }}
            .diff th {{
                background-color: #f0f0f0;
                padding: 5px;
                border: 1px solid #ddd;
                text-align: center;
            }}
            .legend {{
                margin-top: 20px;
                padding: 10px;
                background-color: #f9f9f9;
                border-radius: 5px;
            }}
            .legend ul {{
                padding-left: 20px;
            }}
            .back-button {{
                padding: 8px 15px;
                background-color: #4a6fa5;
                color: white;
                border: none;
                border-radius: 5px;
                cursor: pointer;
                text-decoration: none;
                display: inline-block;
                font-size: 14px;
            }}
            .back-button:hover {{
                background-color: #375d8a;
            }}
            .footer {{
                margin-top: 30px;
                text-align: center;
            }}
        </style>
    </head>
    <body>
        <div class="diff-header">
            <h1>Text Comparison Results</h1>
            <button onclick="goBack()" class="back-button">Back to Report Options</button>
        </div>

        {diff_html}

        <div class="legend">
            <h3>Legend:</h3>
            <ul>
                <li><span style="color: green; font-weight: bold;">Added content</span> - Content that appears in the comparison text but not in the original</li>
                <li><span style="color: red; text-decoration: line-through;">Removed content</span> - Content that appears in the original text but not in the comparison</li>
                <li><span style="color: #996600; font-weight: bold;">Changed content</span> - Content that has been modified between the texts</li>
            </ul>
        </div>

        <div class="footer">
            <button onclick="goBack()" class="back-button">Back to Report Options</button>
        </div>

        <script>
        function goBack() {{
            window.location.href = window.location.origin + '/finalise_report';
        }}
        </script>
    </body>
    </html>
    """

    return styled_html


# Add these routes to your application.py file

@application.route('/report_builder')
@login_required
def enhanced_report_builder():
    """Render the enhanced report builder page"""
    return render_template('report_builder.html')


@application.route('/api/get_rooms')
@login_required
def get_rooms():
    """API endpoint to get rooms from a CSV file"""
    csv_id = request.args.get('csvId')

    if not csv_id:
        return jsonify({
            'status': 'error',
            'message': 'Missing csvId parameter'
        })

    # Find the CSV file
    csv_path = None
    if csv_id == 'latest':
        # Find the most recent CSV file
        transcript_dir = Path('/tmp/temp_transcripts')
        if not transcript_dir.exists():
            return jsonify({
                'status': 'error',
                'message': 'No transcript directory found'
            })

        csv_files = list(transcript_dir.glob('*.csv'))
        if not csv_files:
            return jsonify({
                'status': 'error',
                'message': 'No CSV files found'
            })

        # Sort by modification time, newest first
        csv_files.sort(key=lambda x: x.stat().st_mtime, reverse=True)
        csv_path = csv_files[0]
    else:
        # Find the specific CSV file
        csv_path = UPLOAD_FOLDER / f"{csv_id}.csv"
        if not csv_path.exists():
            return jsonify({
                'status': 'error',
                'message': 'CSV file not found'
            })

    try:
        # Read the CSV file
        df = pd.read_csv(csv_path)

        # Extract unique room names (non-empty ones)
        rooms = []
        current_room = None

        for _, row in df.iterrows():
            room = row['Room'].strip() if pd.notna(row['Room']) and row['Room'].strip() else None
            if room:
                current_room = room
                if current_room not in rooms:
                    rooms.append(current_room)

        return jsonify({
            'status': 'success',
            'rooms': rooms
        })
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': f'Error processing CSV: {str(e)}'
        })


@application.route('/api/prepare_report', methods=['POST'])
@login_required
def prepare_report():
    """API endpoint to prepare a report and get the intermediate page"""
    report_type = request.form.get('reportType', 'full')
    use_latest = request.form.get('useLatest') == 'true'

    try:
        csv_path = None
        csv_id = None

        if use_latest:
            # Find the most recent CSV file
            transcript_dir = Path('/tmp/temp_transcripts')
            if not transcript_dir.exists():
                return jsonify({
                    'status': 'error',
                    'message': 'No transcript directory found'
                })

            csv_files = list(transcript_dir.glob('*.csv'))
            if not csv_files:
                return jsonify({
                    'status': 'error',
                    'message': 'No CSV files found'
                })

            # Sort by modification time, newest first
            csv_files.sort(key=lambda x: x.stat().st_mtime, reverse=True)
            csv_path = csv_files[0]
            csv_id = 'latest'
        else:
            if 'csv' not in request.files:
                return jsonify({
                    'status': 'error',
                    'message': 'No CSV file provided'
                })

            csv_file = request.files['csv']
            if csv_file.filename == '':
                return jsonify({
                    'status': 'error',
                    'message': 'No file selected'
                })

            # Generate a unique ID for this CSV
            csv_id = str(uuid.uuid4())

            # Save uploaded file
            csv_path = UPLOAD_FOLDER / f"{csv_id}.csv"
            csv_file.save(csv_path)

        # Redirect to the report builder with the CSV ID and report type
        # Fixed URL to match your route name
        return jsonify({
            'status': 'success',
            'redirectUrl': f'/report_builder?csvId={csv_id}&type={report_type}',
            'message': 'CSV processed successfully'
        })

    except Exception as e:
        # Clean up temporary file if it exists and was uploaded
        if csv_path and not use_latest and csv_path.exists():
            csv_path.unlink(missing_ok=True)

        return jsonify({
            'status': 'error',
            'message': f'Error processing CSV: {str(e)}'
        })


@application.route('/api/generate_enhanced_report', methods=['POST'])
@login_required
def generate_enhanced_report():
    """API endpoint to generate a final report with property details and room images"""
    try:

        # Add detailed request debugging
        logging.info(f"Request method: {request.method}")
        logging.info(f"Request content type: {request.content_type}")
        logging.info(f"Request content length: {request.content_length}")
        logging.info(f"Request headers: {dict(request.headers)}")

        # Get form data
        report_type = request.form.get('reportType', 'full')
        csv_id = request.form.get('csvId')
        address = request.form.get('address')
        inspection_date = request.form.get('inspectionDate')
        on_behalf_of = request.form.get('onBehalfOf')

        if not csv_id or not address or not inspection_date or not on_behalf_of:
            return jsonify({
                'status': 'error',
                'message': 'Missing required fields'
            })

        # Find the CSV file
        csv_path = None
        if csv_id == 'latest':
            # Find the most recent CSV file
            transcript_dir = Path('/tmp/temp_transcripts')
            if not transcript_dir.exists():
                return jsonify({
                    'status': 'error',
                    'message': 'No transcript directory found'
                })

            csv_files = list(transcript_dir.glob('*.csv'))
            if not csv_files:
                return jsonify({
                    'status': 'error',
                    'message': 'No CSV files found'
                })

            # Sort by modification time, newest first
            csv_files.sort(key=lambda x: x.stat().st_mtime, reverse=True)
            csv_path = csv_files[0]
        else:
            # Find the specific CSV file
            csv_path = UPLOAD_FOLDER / f"{csv_id}.csv"
            if not csv_path.exists():
                return jsonify({
                    'status': 'error',
                    'message': 'CSV file not found'
                })

        # Process uploaded images with simplified field names
        room_images = {}

        # Debug: print all keys in request.form and request.files
        logging.info(f"Form keys: {list(request.form.keys())}")
        logging.info(f"File keys: {list(request.files.keys())}")

        # Process files with simpler naming convention
        for key in request.files.keys():
            if key.startswith('roomFile_'):
                # Extract index information from the key
                parts = key.split('_')
                if len(parts) >= 3:
                    room_index = parts[1]
                    file_index = parts[2]

                    # Get the corresponding room name
                    room_name_key = f"roomName_{room_index}_{file_index}"
                    room_name = request.form.get(room_name_key, '').lower()

                    if room_name:
                        logging.info(f"Processing image for room: {room_name}")

                        # Initialize room in the dictionary if needed
                        if room_name not in room_images:
                            room_images[room_name] = []

                        # Save the file
                        file = request.files[key]
                        if file and file.filename:
                            # Generate a unique filename
                            img_id = str(uuid.uuid4())
                            img_ext = Path(secure_filename(file.filename)).suffix
                            img_path = UPLOAD_FOLDER / f"{img_id}{img_ext}"

                            # Save the image
                            file.save(img_path)
                            room_images[room_name].append(str(img_path))

                            # Log the save
                            logging.info(f"Saved image for {room_name}: {img_path} (from {file.filename})")

        # Also check for the original format (as a fallback)
        for key in request.files.keys():
            if key.startswith('roomImages['):
                # Extract room name from the input name format: roomImages[Room Name]
                room_name = key[11:-1].lower()  # Extract what's between 'roomImages[' and ']'
                logging.info(f"Processing images for room (bracket format): {room_name}")

                # Initialize room in the dictionary if needed
                if room_name not in room_images:
                    room_images[room_name] = []

                # Get all files for this room
                files = request.files.getlist(key)
                logging.info(f"Number of files for {room_name}: {len(files)}")

                # Save each image file
                for file in files:
                    if file and file.filename:
                        # Generate a unique filename
                        img_id = str(uuid.uuid4())
                        img_ext = Path(secure_filename(file.filename)).suffix
                        img_path = UPLOAD_FOLDER / f"{img_id}{img_ext}"

                        # Save the image
                        file.save(img_path)
                        room_images[room_name].append(str(img_path))

                        # Log the save
                        logging.info(f"Saved image for {room_name}: {img_path} (from {file.filename})")

        # Debug log the final image count per room
        for room, images in room_images.items():
            logging.info(f"Room {room}: {len(images)} images collected")

        # Generate the enhanced report
        result_path = generate_enhanced_docx_report(
            csv_path,
            report_type,
            address,
            inspection_date,
            on_behalf_of,
            room_images
        )

        # Create a filename based on address and date
        formatted_date = datetime.strptime(inspection_date, '%Y-%m-%d').strftime('%d-%m-%Y')
        report_name = f"{address.replace(' ', '_')}_{formatted_date}_{report_type}_report.docx"

        return jsonify({
            'status': 'success',
            'reportUrl': f"/download_report/{result_path.name}",
            'reportName': report_name,
            'message': f'{report_type.capitalize()} report generated successfully'
        })

    except Exception as e:
        logging.error(f"Error generating report: {str(e)}", exc_info=True)
        return jsonify({
            'status': 'error',
            'message': f'Error generating report: {str(e)}'
        })


@application.route('/api/get_csv_rooms')
@login_required
def get_csv_rooms():
    """API endpoint to get rooms from a CSV file"""
    csv_id = request.args.get('csvId')

    if not csv_id:
        return jsonify({
            'status': 'error',
            'message': 'Missing csvId parameter'
        })

    # Find the CSV file
    csv_path = None
    if csv_id == 'latest':
        # Find the most recent CSV file
        transcript_dir = Path('/tmp/temp_transcripts')
        if not transcript_dir.exists():
            return jsonify({
                'status': 'error',
                'message': 'No transcript directory found'
            })

        csv_files = list(transcript_dir.glob('*.csv'))
        if not csv_files:
            return jsonify({
                'status': 'error',
                'message': 'No CSV files found'
            })

        # Sort by modification time, newest first
        csv_files.sort(key=lambda x: x.stat().st_mtime, reverse=True)
        csv_path = csv_files[0]
    else:
        # Find the specific CSV file - try both locations
        csv_path_upload = UPLOAD_FOLDER / f"{csv_id}.csv"
        csv_path_transcript = TRANSCRIPT_FOLDER / f"{csv_id}.csv"

        if csv_path_upload.exists():
            csv_path = csv_path_upload
        elif csv_path_transcript.exists():
            csv_path = csv_path_transcript
        else:
            # Try to see if the full filename was passed
            for folder in [UPLOAD_FOLDER, TRANSCRIPT_FOLDER]:
                potential_path = folder / csv_id
                if potential_path.exists():
                    csv_path = potential_path
                    break

            if not csv_path:
                return jsonify({
                    'status': 'error',
                    'message': 'CSV file not found'
                })

    try:
        # Read the CSV file
        df = pd.read_csv(csv_path)

        # Extract unique room names (non-empty ones)
        rooms = []
        current_room = None

        for _, row in df.iterrows():
            room_col = 'Room' if 'Room' in df.columns else None
            if not room_col:
                return jsonify({
                    'status': 'error',
                    'message': 'Room column not found in CSV'
                })

            room = row[room_col].strip() if pd.notna(row[room_col]) and row[room_col].strip() else None
            if room:
                current_room = room
                if current_room not in rooms:
                    rooms.append(current_room)

        # Convert room names to lowercase for consistency
        rooms = [room.lower() for room in rooms]

        return jsonify({
            'status': 'success',
            'rooms': rooms
        })
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': f'Error processing CSV: {str(e)}'
        })

def generate_enhanced_docx_report(csv_path, report_type, address, inspection_date, on_behalf_of, room_images):
    """
    Generate an enhanced report with property details and room images

    Args:
        csv_path: Path to the CSV file
        report_type: Type of report to generate ('inventory', 'full', or 'checkout')
        address: Property address
        inspection_date: Date of inspection
        on_behalf_of: Prepared on behalf of
        room_images: Dictionary mapping room names to lists of image paths

    Returns:
        Path to the generated report
    """
    import pandas as pd
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from datetime import datetime

    # Read CSV data
    df = pd.read_csv(csv_path)

    # Create a new Document
    doc = Document()

    # Set document properties
    doc.core_properties.title = "Property Inspection Report"
    doc.core_properties.author = "PhillyScript"

    # Format the inspection date
    try:
        # Parse the date from the form (YYYY-MM-DD)
        inspection_date_obj = datetime.strptime(inspection_date, '%Y-%m-%d')
        # Format as Month Day, Year
        formatted_date = inspection_date_obj.strftime('%B %d, %Y')
    except:
        # Fallback to the provided string if parsing fails
        formatted_date = inspection_date

    # REPLACE ALL THIS BOILERPLATE CODE:
    # if report_type == 'inventory':
    #     # INVENTORY CHECK-IN REPORT FORMAT
    #     title = doc.add_heading('INVENTORY', 0)
    #     ...
    # WITH THIS SINGLE LINE:
    generate_report_boilerplate(doc, report_type, address, formatted_date, on_behalf_of)

    # Group the data by room - using a simple approach to handle empty room cells
    room_data = {}
    current_room = None

    # First pass - group rows by room
    for _, row in df.iterrows():
        room = row['Room'].strip() if pd.notna(row['Room']) and row['Room'].strip() else None

        if room:
            current_room = room

        if current_room not in room_data:
            room_data[current_room] = []

        room_data[current_room].append(row)

    # Update the room processing in generate_enhanced_docx_report function
    # This restructures how we add room content to ensure images come at the end

    # Second pass - process each room group
    first_room = True
    room_counter = 1  # Counter for room numbering

    for room, rows in room_data.items():
        if room is None:
            continue  # Skip rows with no room (should be rare/nonexistent)

        # Add a page break except for the first room
        if not first_room:
            doc.add_page_break()
        first_room = False

        # Add room heading with number
        numbered_room_heading = f"{room_counter}. {room}"
        room_heading = doc.add_heading(numbered_room_heading, level=1)
        room_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Create a table for this room's inventory items
        # Columns: Attribute, Feature, Comment, Tenant Responsibility (TR)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set the header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Attribute'
        header_cells[1].text = 'Feature'
        header_cells[2].text = 'Comment'
        header_cells[3].text = 'Tenant Responsibility (TR)'

        # Apply header formatting
        for cell in header_cells:
            cell_para = cell.paragraphs[0]
            cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell_para.runs:
                run.bold = True

        # Add the data rows for this room
        # Track attributes to add numbering
        attribute_counter = 1
        current_attribute = None

        for row_data in rows:
            # Get attribute value
            attribute = row_data['Attribute'] if pd.notna(row_data['Attribute']) else ''

            # Update attribute numbering if this is a new attribute
            if attribute and attribute != current_attribute:
                current_attribute = attribute
                numbered_attribute = f"{room_counter}.{attribute_counter} {attribute}"
                attribute_counter += 1
            else:
                # Empty string if this row doesn't have an attribute (continuing from previous)
                numbered_attribute = ''

            # Add a new row to the table
            new_row = table.add_row().cells

            # Fill in the cells
            new_row[0].text = numbered_attribute
            new_row[1].text = row_data['Feature'] if pd.notna(row_data['Feature']) else ''
            new_row[2].text = row_data['Comment'] if pd.notna(row_data['Comment']) else ''
            new_row[3].text = row_data['Tenant Responsibility (TR)'] if pd.notna(
                row_data['Tenant Responsibility (TR)']) else ''

            # Center align the TR column
            if new_row[3].text:
                new_row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # NOW add room images AFTER the inventory table - this ensures they appear at the end of the room section
        # This is the key change from the previous version
        room_lower = room.lower() if isinstance(room, str) else ""
        logging.info(f"Looking for images for room: {room_lower}")
        logging.info(f"Available room keys in room_images: {list(room_images.keys())}")

        if room_lower in room_images and room_images[room_lower]:
            image_count = len(room_images[room_lower])
            logging.info(f"Found {image_count} images for room: {room_lower}")

            # Add images heading if we have images
            if image_count > 0:
                img_heading = doc.add_heading('Room Images', level=2)

            # Add each image in sequence
            for i, img_path in enumerate(room_images[room_lower]):
                logging.info(f"Adding image {i + 1}: {img_path}")
                try:
                    # Create a paragraph for the image
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Verify the image path exists
                    if not os.path.exists(img_path):
                        logging.error(f"Image file not found: {img_path}")
                        p.add_run(f"[Image {i + 1} file not found]")
                        continue

                    # Log image file details
                    file_size = os.path.getsize(img_path)
                    logging.info(f"Image file size: {file_size} bytes")

                    # Add the image to the paragraph
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(5.0))  # Standard width

                    # Add caption
                    caption = doc.add_paragraph(f"Image {i + 1}")
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.style = 'Caption'

                    # Add some space
                    doc.add_paragraph()

                    logging.info(f"Successfully added image {i + 1}")

                except Exception as e:
                    # If adding the image fails, add a placeholder text
                    error_p = doc.add_paragraph(f"[Image {i + 1} could not be displayed: {str(e)}]")
                    error_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    logging.error(f"Error adding image to report: {str(e)}", exc_info=True)

                    # Add space even if the image fails
                    doc.add_paragraph()
        else:
            logging.info(f"No images found for room: {room_lower}")

        # Increment room counter for the next room
        room_counter += 1

    # REPLACE ALL THIS CLOSING CODE:
    # # Add a page break before closing sections
    # doc.add_page_break()
    #
    # # Add closing content based on report type
    # if report_type == 'inventory':
    #     # Add tenant and landlord declaration section
    #     ...
    # WITH THIS SINGLE LINE:
    generate_report_closing(doc, report_type, 124)  # 124 is a placeholder for page count

    # Save the document with report type in the filename
    result_id = str(uuid.uuid4())
    filename = f"{result_id}_{report_type}_report.docx"
    result_path = RESULT_FOLDER / filename
    doc.save(result_path)

    # Clean up temporary image files
    for room_img_list in room_images.values():
        for img_path in room_img_list:
            try:
                Path(img_path).unlink(missing_ok=True)
            except:
                pass

    return result_path

'''Admin Functions'''

@application.route('/network_test')
def network_test():
    """Comprehensive network diagnostic tool"""
    import socket
    import requests
    import subprocess
    import traceback
    import json

    results = {
        "timestamp": str(datetime.now()),
        "environment": {},
        "dns_tests": {},
        "connection_tests": {},
        "network_info": {},
        "trace_routes": {}
    }

    # Get environment info
    try:
        results["environment"] = {
            "platform": platform.platform(),
            "python_version": platform.python_version(),
            "hostname": socket.gethostname()
        }

        # Get environment variables (redact sensitive ones)
        env_vars = {}
        for key, value in os.environ.items():
            if any(sensitive in key.lower() for sensitive in ['key', 'secret', 'pass', 'token']):
                env_vars[key] = "REDACTED"
            else:
                env_vars[key] = value
        results["environment"]["env_vars"] = env_vars
    except Exception as e:
        results["environment"]["error"] = str(e)

    # DNS lookup tests
    domains = [
        'speech.googleapis.com',
        'www.google.com',
        'transcribe.amazonaws.com',
        's3.amazonaws.com',
        'ec2.amazonaws.com',
        'api.openai.com',
        'phillyscript-db-1.ct2ce24yc54l.eu-west-2.rds.amazonaws.com'  # Your RDB
    ]

    for domain in domains:
        try:
            ip = socket.gethostbyname(domain)
            results["dns_tests"][domain] = {"status": "success", "ip": ip}
        except Exception as e:
            results["dns_tests"][domain] = {"status": "failed", "error": str(e)}

    # Connection tests
    for domain in domains:
        if results["dns_tests"].get(domain, {}).get("status") == "success":
            try:
                requests_timeout = 3
                url = f"https://{domain}"

                # Use a very short timeout for quicker failure
                response = requests.get(url, timeout=requests_timeout)
                results["connection_tests"][domain] = {
                    "status": "success",
                    "status_code": response.status_code,
                    "reason": response.reason,
                    "timeout_used": requests_timeout
                }
            except requests.exceptions.ConnectTimeout:
                results["connection_tests"][domain] = {
                    "status": "failed",
                    "error": "Connection timed out",
                    "timeout_used": requests_timeout
                }
            except requests.exceptions.ConnectionError as e:
                results["connection_tests"][domain] = {
                    "status": "failed",
                    "error": f"Connection error: {str(e)}"
                }
            except Exception as e:
                results["connection_tests"][domain] = {
                    "status": "failed",
                    "error": str(e),
                    "traceback": traceback.format_exc()
                }

    # Get network interface info
    try:
        results["network_info"]["interfaces"] = {}

        # Get network interfaces and IPs
        import netifaces
        for iface in netifaces.interfaces():
            addrs = netifaces.ifaddresses(iface)
            if netifaces.AF_INET in addrs:  # IPv4 info
                results["network_info"]["interfaces"][iface] = addrs[netifaces.AF_INET]
    except ImportError:
        results["network_info"]["interfaces"] = "netifaces module not available"
    except Exception as e:
        results["network_info"]["interfaces"] = {"error": str(e)}

    # Try to get routing information
    try:
        routes_output = subprocess.check_output(["ip", "route"], stderr=subprocess.STDOUT, timeout=3).decode('utf-8')
        results["network_info"]["routes"] = routes_output.split('\n')
    except (subprocess.SubprocessError, FileNotFoundError):
        try:
            # Alternative method
            routes_output = subprocess.check_output(["netstat", "-rn"], stderr=subprocess.STDOUT, timeout=3).decode(
                'utf-8')
            results["network_info"]["routes"] = routes_output.split('\n')
        except:
            results["network_info"]["routes"] = "Could not retrieve route information"

    # IMPORTANT: Check if proxy settings are affecting connections
    try:
        proxy_info = {
            "http_proxy": os.environ.get("http_proxy", "Not set"),
            "https_proxy": os.environ.get("https_proxy", "Not set"),
            "HTTP_PROXY": os.environ.get("HTTP_PROXY", "Not set"),
            "HTTPS_PROXY": os.environ.get("HTTPS_PROXY", "Not set"),
            "no_proxy": os.environ.get("no_proxy", "Not set"),
            "NO_PROXY": os.environ.get("NO_PROXY", "Not set")
        }
        results["network_info"]["proxy_settings"] = proxy_info
    except Exception as e:
        results["network_info"]["proxy_settings"] = {"error": str(e)}

    # Try ping/traceroute (these often don't work in containers but worth trying)
    for domain in ['speech.googleapis.com', 'transcribe.amazonaws.com']:
        try:
            ping_output = subprocess.check_output(["ping", "-c", "3", domain], stderr=subprocess.STDOUT,
                                                  timeout=5).decode('utf-8')
            results["trace_routes"][domain] = {"ping": ping_output.split('\n')}
        except (subprocess.SubprocessError, FileNotFoundError):
            results["trace_routes"][domain] = {"ping": "Failed to ping"}

    # Return formatted JSON
    return jsonify(results)


@application.route('/test_openai')
def test_openai():
    import openai
    import json

    results = {"status": "unknown", "details": {}}

    try:
        # Same client initialization as your production code
        client = openai.OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

        # Simplest possible API call
        response = client.models.list()

        # Return success details
        results["status"] = "success"
        results["details"] = {
            "models": [model.id for model in response.data[:5]],  # Just show first 5 models
            "total_models": len(response.data)
        }
    except Exception as e:
        # Capture detailed error information
        results["status"] = "error"
        results["details"] = {
            "error_type": type(e).__name__,
            "error_message": str(e),
            "traceback": traceback.format_exc()
        }

    return jsonify(results)



def transcribe_with_dummy_data(original_filename):
    """Return a dummy transcript with the correct structure as a fallback"""
    print(f"Using dummy transcript data for {original_filename}")

    # Create predetermined structured transcript
    transcript = """
new room living room and kitchen
new attribute doors
new feature front door
comment in good condition
comment hindu religious symbol
new feature lock on door
comment in good condition
comment light scuff marks
new feature cupboard door
comment scratches around door handle
new feature floor by doorway
comment stain marks around door frames
new feature shoe rack
comment handle loose on lower drawer
new attribute skirting boards
new feature all skirting boards
comment in generally good condition
"""

    return transcript.strip()


@application.route('/test_google_speech')
def test_google_speech():
    """Direct test of Google Speech API"""
    import speech_recognition as sr
    import requests
    import socket
    import time

    results = {
        "timestamp": str(datetime.now()),
        "tests": []
    }

    # Test 1: Basic DNS resolution
    test1 = {"name": "DNS Resolution", "target": "speech.googleapis.com"}
    try:
        start_time = time.time()
        ip = socket.gethostbyname("speech.googleapis.com")
        test1["status"] = "success"
        test1["ip"] = ip
        test1["time_taken"] = f"{(time.time() - start_time):.2f}s"
    except Exception as e:
        test1["status"] = "failed"
        test1["error"] = str(e)
    results["tests"].append(test1)

    # Test 2: HTTPS connection
    test2 = {"name": "HTTPS Connection", "target": "https://speech.googleapis.com/v1/speech:recognize"}
    try:
        start_time = time.time()
        response = requests.get("https://speech.googleapis.com/v1/speech:recognize", timeout=5)
        test2["status"] = "success"
        test2["status_code"] = response.status_code
        test2["response"] = response.text[:100] + "..." if len(response.text) > 100 else response.text
        test2["time_taken"] = f"{(time.time() - start_time):.2f}s"
    except Exception as e:
        test2["status"] = "failed"
        test2["error"] = str(e)
    results["tests"].append(test2)

    # Test 3: Speech Recognizer initialization
    test3 = {"name": "Speech Recognizer Init"}
    try:
        start_time = time.time()
        recognizer = sr.Recognizer()
        test3["status"] = "success"
        test3["time_taken"] = f"{(time.time() - start_time):.2f}s"
    except Exception as e:
        test3["status"] = "failed"
        test3["error"] = str(e)
    results["tests"].append(test3)

    # Only do further tests if recognizer initialized successfully
    if test3["status"] == "success":
        # Create a small audio file in-memory to test
        try:
            from pydub import AudioSegment
            from pydub.generators import Sine

            # Generate a simple tone
            test4 = {"name": "Generate Test Audio"}
            start_time = time.time()
            tone = Sine(440).to_audio_segment(duration=1000)  # 1 second 440Hz tone
            test_audio_path = "/tmp/test_tone.wav"
            tone.export(test_audio_path, format="wav")
            test4["status"] = "success"
            test4["time_taken"] = f"{(time.time() - start_time):.2f}s"
            results["tests"].append(test4)

            # Try to recognize with Google (expecting failure, but want to see exact error)
            test5 = {"name": "Google API Recognition"}
            try:
                start_time = time.time()
                with sr.AudioFile(test_audio_path) as source:
                    audio_data = recognizer.record(source)
                    # This will likely fail, but we want to see the exact error
                    transcript = recognizer.recognize_google(audio_data)
                    test5["status"] = "success"  # Surprising if it works!
                    test5["transcript"] = transcript
                    test5["time_taken"] = f"{(time.time() - start_time):.2f}s"
            except Exception as e:
                test5["status"] = "failed"
                test5["error"] = str(e)
                test5["error_type"] = type(e).__name__
            results["tests"].append(test5)

        except Exception as e:
            results["tests"].append({
                "name": "Generate Test Audio",
                "status": "failed",
                "error": str(e)
            })

    return jsonify(results)

@app.route('/test_outbound')
def test_outbound():
    import requests
    try:
        response = requests.get('https://speech.googleapis.com/v1/speech', timeout=5)
        return f"Outbound connectivity test: {response.status_code}"
    except Exception as e:
        return f"Outbound connectivity test failed: {str(e)}"

@application.route('/test_google')
def test_google():
    import urllib.request
    import socket
    try:
        socket.setdefaulttimeout(5)
        response = urllib.request.urlopen('https://www.google.com')
        return f"Connection successful! Status code: {response.getcode()}"
    except Exception as e:
        return f"Connection failed: {str(e)}"

@app.route('/init_db', methods=['GET'])
def init_db():
    try:
        # Create admin user if it doesn't exist
        admin_user = User.query.filter_by(username='admin').first()
        if not admin_user:
            admin_user = User(
                username='admin',
                email='admin@example.com',
                role='master_admin'
            )
            admin_user.set_password('changeme123')
            db.session.add(admin_user)
            db.session.commit()
            return 'Admin user created successfully!'
        else:
            return 'Admin user already exists!'
    except Exception as e:
        return f'Error: {str(e)}'

@application.route('/debug_environment')
@login_required
def debug_environment():
    """Debug endpoint to show environment details"""

    env_info = {
        'python_version': sys.version,
        'platform': platform.platform(),
        'environment_vars': dict(os.environ),
        'upload_folder': str(UPLOAD_FOLDER),
        'upload_folder_exists': os.path.exists(UPLOAD_FOLDER),
        'temp_dir': tempfile.gettempdir(),
        'disk_space': shutil.disk_usage('/'),
    }

    return jsonify(env_info)


def test_database_connectivity():
    """
    Comprehensive database connectivity test with detailed diagnostics.

    Returns:
        dict: Connectivity test results with detailed information
    """
    results = {
        "status": "unknown",
        "connection": {},
        "engine_details": {},
        "environment": {}
    }

    try:
        # Retrieve database URL from environment
        database_url = os.environ.get('DATABASE_URL')

        if not database_url:
            return {
                "status": "error",
                "message": "DATABASE_URL environment variable not set"
            }

        # Logging configuration
        logging.basicConfig(level=logging.INFO)
        logger = logging.getLogger(__name__)

        # Create engine with extended diagnostics
        engine = create_engine(
            database_url,
            echo=True,  # Enable SQL logging
            pool_pre_ping=True,  # Test connection before using
            pool_recycle=3600,  # Recycle connections after 1 hour
            connect_args={
                'connect_timeout': 10,  # 10-second connection timeout
            }
        )

        # Collect comprehensive environment details
        results['environment'] = {
            'database_url': make_url(database_url).render_as_string(hide_password=True),
            'python_version': platform.python_version(),
            'python_implementation': platform.python_implementation(),
            'system': platform.system(),
            'system_release': platform.release(),
            'sqlalchemy_version': sqlalchemy.__version__,
            'python_executable': sys.executable
        }

        # Test connection
        with engine.connect() as connection:
            # Simple query to test connectivity
            result = connection.execute(text("SELECT current_timestamp"))
            current_time = result.fetchone()[0]

            # Get connection details using psycopg2-specific methods
            raw_connection = connection.connection

            # Attempt to get database server version
            try:
                version_query = connection.execute(text("SELECT version()"))
                db_version = version_query.scalar()
            except Exception as version_err:
                db_version = "Version query failed: " + str(version_err)

            results['connection'] = {
                'status': 'success',
                'current_database_time': str(current_time),
                'connection_pool_status': str(engine.pool.status())
            }

            # Engine and connection details
            results['engine_details'] = {
                'database_server_version': db_version,
                'connection_parameters': {
                    'dsn': raw_connection.dsn if hasattr(raw_connection, 'dsn') else 'Not available',
                    'autocommit': raw_connection.autocommit if hasattr(raw_connection,
                                                                       'autocommit') else 'Not available'
                }
            }

        results['status'] = 'success'

    except SQLAlchemyError as e:
        # Specific SQLAlchemy error handling
        results['status'] = 'error'
        results['connection']['error'] = {
            'type': type(e).__name__,
            'message': str(e),
            'detailed_traceback': traceback.format_exc()
        }

        # Log the full error for server-side tracking
        logger.error(f"Database Connectivity Test Failed: {e}")
        logger.error(traceback.format_exc())

    except Exception as e:
        # Catch-all for any other unexpected errors
        results['status'] = 'error'
        results['connection']['error'] = {
            'type': type(e).__name__,
            'message': str(e),
            'detailed_traceback': traceback.format_exc()
        }

        # Log the full error for server-side tracking
        logger.error(f"Unexpected Error in Database Connectivity Test: {e}")
        logger.error(traceback.format_exc())

    return results


@application.route('/test_database', methods=['GET'])
def database_connectivity_route():
    """
    Web route to test database connectivity.
    Exposes the database test function via a Flask route.
    """
    test_results = test_database_connectivity()
    return jsonify(test_results)

@application.route('/test_upload_page')
@login_required
def test_upload_page():
    """Render the test upload page"""
    return render_template('test_upload.html')


@application.route('/test_file_upload', methods=['POST'])
def test_file_upload():
    """Test endpoint for file uploads"""
    logging.info(f"Test upload - Content type: {request.content_type}")
    logging.info(f"Test upload - Content length: {request.content_length}")
    logging.info(f"Test upload - Form keys: {list(request.form.keys())}")
    logging.info(f"Test upload - File keys: {list(request.files.keys())}")

    file_details = []
    for key in request.files:
        file = request.files[key]
        if file.filename:
            # Save to test location
            path = Path(UPLOAD_FOLDER) / f"test_{secure_filename(file.filename)}"
            file.save(path)
            size = path.stat().st_size
            file_details.append({
                'key': key,
                'filename': file.filename,
                'size': size
            })
            logging.info(f"Test saved file: {path}, size: {size}")

    return jsonify({
        'status': 'received',
        'form_keys': list(request.form.keys()),
        'file_keys': list(request.files.keys()),
        'files': file_details
    })

@application.errorhandler(Exception)
def handle_error(e):
    logger.error(f"Unhandled exception: {str(e)}")
    return jsonify({
        'error': 'Internal server error',
        'message': str(e)
    }), 500

if __name__ == '__main__':
    # Determine optimal number of threads
    num_threads = min(multiprocessing.cpu_count() * 2, 8)  # Use 2x CPU cores, max 8

    # Get port from environment or use 8080 as default for App Runner
    port = int(os.environ.get('PORT', 8080))

    print(f"Starting PhillyScript server with {num_threads} worker threads on port {port}...")
    print("Available routes:")
    print(f"  - http://127.0.0.1:{port}/ (main interface)")
    print(f"  - http://127.0.0.1:{port}/transcribe (transcription page)")
    print(f"  - http://127.0.0.1:{port}/diff_check (image comparison)")

    # Set threaded mode and optimal worker count
    application.run(debug=False, host='0.0.0.0', port=port, threaded=True)