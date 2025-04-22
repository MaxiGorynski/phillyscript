from flask import Flask, request, render_template, send_file, jsonify, send_from_directory
import os
import uuid
from pathlib import Path
import logging
import boto3
from io import BytesIO
import logging
import shutil
import openai
import time
import tempfile
import subprocess

# Configure logging first - increase level to see more details
logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Initialise Flask app
application = Flask(__name__)
app = application  # AWS ElasticBeanstalk looks for 'application' while Flask CLI looks for 'app'

# Configure app before importing other modules
application.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev-key-for-local-only')

client = openai.OpenAI(
    api_key=os.environ.get("OPENAI_API_KEY")
)

# Get DATABASE_URL from environment, with a SQLite fallback
database_url = os.environ.get('DATABASE_URL')
logger.info(f"Database URL from environment: {database_url}")

# Determine which database to use
if not database_url or 'rds.amazonaws.com' in database_url:
    fallback_db = 'sqlite:///fallback.db'
    logger.warning(f"Using fallback database: {fallback_db}")
    application.config['SQLALCHEMY_DATABASE_URI'] = fallback_db
else:
    application.config['SQLALCHEMY_DATABASE_URI'] = database_url

# Now that the database URI is set, we can try to download from S3 if it's SQLite
if application.config['SQLALCHEMY_DATABASE_URI'].startswith('sqlite'):
    db_file = application.config['SQLALCHEMY_DATABASE_URI'].replace('sqlite:///', '')
    s3_bucket = os.environ.get('S3_BUCKET', 'your-backup-bucket-name')

    try:
        # Download DB from S3 if it exists
        s3 = boto3.client('s3')
        s3.download_file(s3_bucket, 'db_backup/' + os.path.basename(db_file), db_file)
        logger.info(f"Downloaded database from S3: {db_file}")
    except Exception as e:
        logger.warning(f"Could not download database from S3: {str(e)}")

# Define backup function to be used later
def backup_db_to_s3():
    if application.config['SQLALCHEMY_DATABASE_URI'].startswith('sqlite'):
        db_file = application.config['SQLALCHEMY_DATABASE_URI'].replace('sqlite:///', '')
        s3_bucket = os.environ.get('S3_BUCKET', 'your-backup-bucket-name')

        try:
            # Upload DB to S3
            s3 = boto3.client('s3')
            s3.upload_file(db_file, s3_bucket, 'db_backup/' + os.path.basename(db_file))
            logger.info(f"Backed up database to S3: {db_file}")
        except Exception as e:
            logger.warning(f"Could not back up database to S3: {str(e)}")

application.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
application.config['TEMPLATES_AUTO_RELOAD'] = True

# Database configuration needs to be engine-specific
if application.config['SQLALCHEMY_DATABASE_URI'].startswith('sqlite'):
    # SQLite-specific options
    application.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
        'connect_args': {'check_same_thread': False}
    }
else:
    # PostgreSQL-specific options
    application.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
        'pool_recycle': 280,
        'pool_timeout': 10,
        'connect_args': {
            'connect_timeout': 5  # PostgreSQL connection timeout in seconds
        }
    }

# Log the database we're connecting to
logger.info(f"Configured database: {application.config['SQLALCHEMY_DATABASE_URI']}")

S3_AVAILABLE = False

try:
    import boto3
    S3_AVAILABLE = True
except ImportError:
    logger.warning("boto3 not available, S3 integration will be disabled")


def backup_db_to_s3():
    if not S3_AVAILABLE:
        logger.warning("S3 not available, skipping backup")
        return False

    if application.config['SQLALCHEMY_DATABASE_URI'].startswith('sqlite'):
        db_file = application.config['SQLALCHEMY_DATABASE_URI'].replace('sqlite:///', '')
        s3_bucket = os.environ.get('S3_BUCKET')

        if not s3_bucket:
            logger.warning("S3_BUCKET not set, skipping backup")
            return False

        try:
            # Upload DB to S3
            s3 = boto3.client('s3')
            s3.upload_file(db_file, s3_bucket, 'db_backup/' + os.path.basename(db_file))
            logger.info(f"Backed up database to S3: {db_file}")
            return True
        except Exception as e:
            logger.warning(f"Could not back up database to S3: {str(e)}")
            return False
    return False

# Make backup function available globally
application.backup_db_to_s3 = backup_db_to_s3

# Import db and login_manager from extensions
from extensions import db, login_manager

# Initialize extensions with the application
db.init_app(application)
login_manager.init_app(application)
login_manager.login_view = 'auth.login'

# Create folders for file storage
UPLOAD_FOLDER = Path('/tmp/temp_uploads')
TRANSCRIPT_FOLDER = Path('/tmp/temp_transcripts')
RESULT_FOLDER = Path('/tmp/results')

for folder in [UPLOAD_FOLDER, TRANSCRIPT_FOLDER, RESULT_FOLDER]:
    folder.mkdir(exist_ok=True, parents=True)


# Create a diagnostic route BEFORE importing models
@application.route('/api/diagnostics')
def diagnostics():
    db_status = "unknown"
    try:
        with db.engine.connect() as conn:
            conn.execute("SELECT 1")
            db_status = "connected"
    except Exception as e:
        db_status = f"error: {str(e)}"

    return jsonify({
        'status': 'ok',
        'message': 'API is running',
        'database': {
            'status': db_status,
            'uri': application.config['SQLALCHEMY_DATABASE_URI'].split('@')[-1].split('/')[0]
            if '@' in application.config['SQLALCHEMY_DATABASE_URI'] else 'local'
        },
        'environment': {
            'FLASK_ENV': os.environ.get('FLASK_ENV', 'production')
        }
    })


# Import models and auth AFTER initializing the database
try:
    from models import User

    logger.info("Models imported successfully")
except Exception as e:
    logger.error(f"Error importing models: {str(e)}")


    # Create a minimal User model for the app to run
    class User(db.Model):
        id = db.Column(db.Integer, primary_key=True)
        username = db.Column(db.String(80), unique=True, nullable=False)
        email = db.Column(db.String(120), unique=True, nullable=False)
        is_active = db.Column(db.Boolean, default=True)

        def get_id(self):
            return str(self.id)

        @property
        def is_authenticated(self):
            return True

        @property
        def is_anonymous(self):
            return False

# Import auth blueprint
try:
    from auth import auth as auth_blueprint

    application.register_blueprint(auth_blueprint)
    logger.info("Auth blueprint registered successfully")
except Exception as e:
    logger.error(f"Error registering auth blueprint: {str(e)}")
    # Create a minimal auth blueprint for the app to run
    from flask import Blueprint

    auth_blueprint = Blueprint('auth', __name__)
    application.register_blueprint(auth_blueprint)


# Configure user loader
@login_manager.user_loader
def load_user(user_id):
    try:
        return User.query.get(int(user_id))
    except:
        return None


# Import the rest of your dependencies
import speech_recognition as sr
from pydub import AudioSegment
import pandas as pd
import cv2
import numpy as np
import re
from werkzeug.utils import secure_filename
import threading
from concurrent.futures import ThreadPoolExecutor
import multiprocessing
import difflib
from docx import Document
from datetime import datetime
from flask_login import login_required, current_user


# Create database tables ONCE within an application context with improved error handling
try:
    with application.app_context():
        # Test the database connection first
        try:
            logger.info("Testing database connection...")
            with db.engine.connect() as conn:
                conn.execute("SELECT 1")
            logger.info("Database connection successful!")

            # Create tables if the connection was successful
            logger.info("Creating database tables...")
            db.create_all()
            logger.info("Database tables created successfully")
        except Exception as e:
            logger.error(f"Error with database: {str(e)}")
            logger.warning("Application will continue with limited functionality")
except Exception as e:
    logger.error(f"Database initialization error: {str(e)}")
    logger.warning("Application will run in fallback mode")

def convert_to_wav(audio_path):
    """
    Convert audio file to WAV format if it's not already WAV.
    Returns path to WAV file.
    """
    print(f"Converting {audio_path} to WAV format...")
    audio_path = Path(audio_path)
    if audio_path.suffix.lower() == '.wav':
        return str(audio_path)

    # Create a temporary WAV file
    wav_path = audio_path.with_suffix('.wav')

    # Convert audio to WAV based on file extension
    file_ext = audio_path.suffix.lower()

    try:
        if file_ext == '.mp3':
            audio = AudioSegment.from_mp3(str(audio_path))
        elif file_ext == '.flac':
            audio = AudioSegment.from_file(str(audio_path), format="flac")
        elif file_ext == '.m4a':
            audio = AudioSegment.from_file(str(audio_path), format="m4a")
        elif file_ext == '.aac':
            audio = AudioSegment.from_file(str(audio_path), format="aac")
        else:
            # Try to automatically detect format
            audio = AudioSegment.from_file(str(audio_path))

        audio.export(wav_path, format="wav")
        print(f"Successfully converted to {wav_path}")
        return str(wav_path)
    except Exception as e:
        print(f"Error converting {audio_path}: {str(e)}")
        # If conversion fails, try to return the original path
        # Speech recognition might still work for some formats
        return str(audio_path)


def optimise_audio_file(file_path, output_dir=None, quality='speech', min_silence_len=700):
    """
    Optimise audio file for transcription by applying multiple compression techniques.

    Args:
        file_path: Path to the original audio file
        output_dir: Directory to save the optimised file (defaults to same directory)
        quality: Optimization profile ('speech' or 'high')
        min_silence_len: Minimum silence length in ms to detect and compress

    Returns:
        Tuple of (path to optimised file, compression stats dictionary)
    """
    start_time = time.time()
    file_path = Path(file_path)

    if output_dir is None:
        output_dir = file_path.parent
    else:
        output_dir = Path(output_dir)
        output_dir.mkdir(exist_ok=True, parents=True)

    # Generate output filename
    output_filename = f"optimised_{file_path.stem}.mp3"
    output_path = output_dir / output_filename

    # Log start of optimization
    logging.info(f"Starting audio optimization for: {file_path}")
    original_size = file_path.stat().st_size
    logging.info(f"Original file size: {original_size / 1024 / 1024:.2f} MB")

    try:
        # Load audio file
        audio = AudioSegment.from_file(str(file_path))
        original_duration = len(audio) / 1000  # in seconds

        # Step 1: Convert to mono if stereo
        was_stereo = audio.channels > 1
        if was_stereo:
            audio = audio.set_channels(1)
            logging.info("Converted audio to mono")

        # Step 2: Downsample to 16kHz for speech (22kHz for high quality)
        target_rate = 16000 if quality == 'speech' else 22050
        if audio.frame_rate != target_rate:
            audio = audio.set_frame_rate(target_rate)
            logging.info(f"Downsampled audio to {target_rate}Hz")

        # Step 3: Detect and compress silence
        if min_silence_len > 0:
            # Detect silence
            silence_threshold = -40  # dB, adjust as needed
            chunks = detect_silence_chunks(audio, silence_threshold, min_silence_len)

            if chunks:
                # Compress silence by keeping just enough to maintain natural pauses
                compressed_silence_len = 300  # ms for each silence
                audio = compress_silence(audio, chunks, compressed_silence_len)
                logging.info(f"Compressed {len(chunks)} silence segments")

        # Step 4: Export with optimised encoding
        # For speech, use very aggressive compression
        if quality == 'speech':
            bitrate = "24k"  # Very low bitrate suitable for speech
        else:
            bitrate = "64k"  # Higher quality for general audio

        # Export to WAV with specified parameters
        audio.export(
            output_path,
            format="mp3",
            parameters=[
                "-ar", str(target_rate),
                "-ac", "1",
                "-b:a", bitrate
            ]
        )

        # Calculate compression stats
        end_time = time.time()
        compressed_size = output_path.stat().st_size
        size_reduction = 1 - (compressed_size / original_size)
        compressed_duration = len(audio) / 1000  # in seconds
        duration_reduction = 1 - (compressed_duration / original_duration)

        # Log results
        logging.info(f"Optimised file saved to: {output_path}")
        logging.info(f"Optimised file size: {compressed_size / 1024 / 1024:.2f} MB")
        logging.info(f"Size reduction: {size_reduction:.2%}")
        logging.info(f"Duration reduction: {duration_reduction:.2%}")
        logging.info(f"Optimization completed in {end_time - start_time:.2f} seconds")

        # Return stats dictionary
        stats = {
            "original_path": str(file_path),
            "optimised_path": str(output_path),
            "original_size": original_size,
            "optimised_size": compressed_size,
            "size_reduction_percent": size_reduction * 100,
            "original_duration": original_duration,
            "optimised_duration": compressed_duration,
            "duration_reduction_percent": duration_reduction * 100,
            "processing_time": end_time - start_time,
            "was_stereo": was_stereo,
            "target_sample_rate": target_rate,
            "silence_segments_compressed": len(chunks) if min_silence_len > 0 else 0
        }

        return output_path, stats

    except Exception as e:
        logging.error(f"Error optimizing audio: {str(e)}")
        return None, {"error": str(e)}


def detect_silence_chunks(audio, silence_threshold=-40, min_silence_len=700):
    """
    Detect silent segments in audio.

    Args:
        audio: AudioSegment to analyze
        silence_threshold: dB threshold below which is considered silence
        min_silence_len: minimum silence length in ms

    Returns:
        List of (start_ms, end_ms) tuples of silence chunks
    """
    # Get audio data as numpy array
    samples = np.array(audio.get_array_of_samples())
    sample_rate = audio.frame_rate

    # Convert to float between -1 and 1
    if audio.sample_width == 2:  # 16-bit
        max_value = 32768.0
    elif audio.sample_width == 1:  # 8-bit
        max_value = 128.0
    else:  # 24 or 32-bit
        max_value = 2 ** (8 * audio.sample_width - 1)

    samples = samples / max_value

    # Calculate dB values (using small epsilon to avoid log(0))
    epsilon = 1e-10
    db_values = 20 * np.log10(np.maximum(np.abs(samples), epsilon))

    # Find segments below threshold
    is_silence = db_values < silence_threshold

    # Convert to milliseconds positions
    silence_chunks = []
    in_silence = False
    silence_start = 0

    # Factor to convert sample position to ms
    ms_per_sample = 1000.0 / sample_rate

    for i, silent in enumerate(is_silence):
        position_ms = i * ms_per_sample

        if silent and not in_silence:
            # Start of silence
            in_silence = True
            silence_start = position_ms
        elif not silent and in_silence:
            # End of silence
            silence_end = position_ms
            silence_duration = silence_end - silence_start

            # Only keep if long enough
            if silence_duration >= min_silence_len:
                silence_chunks.append((int(silence_start), int(silence_end)))

            in_silence = False

    # Check if file ends in silence
    if in_silence:
        silence_end = len(samples) * ms_per_sample
        silence_duration = silence_end - silence_start

        if silence_duration >= min_silence_len:
            silence_chunks.append((int(silence_start), int(silence_end)))

    return silence_chunks


def compress_silence(audio, silence_chunks, target_silence_len=300):
    """
    Compress silence segments in audio to a target length.

    Args:
        audio: AudioSegment to modify
        silence_chunks: List of (start_ms, end_ms) tuples of silence
        target_silence_len: Target length in ms for each silence segment

    Returns:
        New AudioSegment with compressed silence
    """
    if not silence_chunks:
        return audio

    # Sort chunks by start time
    silence_chunks.sort(key=lambda x: x[0])

    # Initialize new audio segments list
    segments = []
    last_end = 0

    for start, end in silence_chunks:
        # Add audio before silence
        if start > last_end:
            segments.append(audio[last_end:start])

        # Add compressed silence - keep the middle portion
        silence_duration = end - start
        if silence_duration > target_silence_len:
            # Calculate start and end points for the portion to keep
            keep_start = start + (silence_duration - target_silence_len) // 2
            keep_end = keep_start + target_silence_len
            segments.append(audio[keep_start:keep_end])
        else:
            # If silence is already shorter than target, keep it all
            segments.append(audio[start:end])

        last_end = end

    # Add remaining audio after last silence
    if last_end < len(audio):
        segments.append(audio[last_end:])

    # Concatenate all segments
    if segments:
        return sum(segments)
    return audio


def process_uploaded_audio(file_path, original_filename):
    """
    Process and optimize an uploaded audio file.

    Args:
        file_path: Path to the uploaded audio file
        original_filename: Original filename

    Returns:
        Tuple of (optimized_path, error_message)
    """
    # Create a separate folder for optimized files
    optimised_folder = Path('/tmp/optimized_audio')
    optimised_folder.mkdir(exist_ok=True, parents=True)

    # Log the process
    logging.info(f"Processing uploaded file: {original_filename}")

    try:
        # Optimize the audio
        optimised_path, stats = optimise_audio_file(
            file_path,
            output_dir=optimised_folder,
            quality='speech',  # Optimize for speech
            min_silence_len=700  # Min silence length to compress (ms)
        )

        if optimised_path is None:
            return None, f"Failed to optimize audio: {stats.get('error', 'Unknown error')}"

        # Log the optimization results
        logging.info(
            f"Audio optimization for {original_filename}: {stats['size_reduction_percent']:.2f}% size reduction")

        # Return the optimized file path
        return optimised_path, None

    except Exception as e:
        logging.error(f"Error processing audio {original_filename}: {str(e)}")
        return None, f"Error processing audio: {str(e)}"

def transcribe_audio_optimized(audio_path):
    """
    Optimized version of transcribe_audio function with improved file path handling.
    """
    try:
        print(f"Converting {audio_path} to WAV format...")
        audio_path = Path(audio_path)

        # Verify file exists
        if not audio_path.exists():
            print(f"Error: File does not exist: {audio_path}")
            return ""

        # Skip conversion for WAV files
        if audio_path.suffix.lower() == '.wav':
            wav_path = str(audio_path)
        else:
            # Create a temporary WAV file with optimized parameters
            wav_path = audio_path.with_suffix('.wav')

            # Convert audio to WAV with optimized parameters
            try:
                if audio_path.suffix.lower() == '.mp3':
                    audio = AudioSegment.from_mp3(str(audio_path))
                elif audio_path.suffix.lower() == '.flac':
                    audio = AudioSegment.from_file(str(audio_path), format="flac")
                elif audio_path.suffix.lower() == '.m4a':
                    audio = AudioSegment.from_file(str(audio_path), format="m4a")
                else:
                    audio = AudioSegment.from_file(str(audio_path))

                # Optimize audio for speech recognition
                # Downsampling to 16kHz, mono, 16-bit which is optimal for most speech recognition
                audio = audio.set_frame_rate(16000).set_channels(1)

                # Export with optimized settings
                audio.export(wav_path, format="wav", parameters=["-q:a", "0"])
                print(f"Successfully converted to {wav_path}")
            except Exception as e:
                print(f"Error converting audio: {e}")
                return ""

        # Verify WAV file exists and has content
        wav_path_obj = Path(wav_path)
        if not wav_path_obj.exists():
            print(f"Error: WAV file does not exist: {wav_path}")
            return ""

        if wav_path_obj.stat().st_size == 0:
            print(f"Error: WAV file is empty: {wav_path}")
            return ""

        print("Starting transcription...")
        recognizer = sr.Recognizer()

        # Set recognition parameters for better performance
        recognizer.energy_threshold = 300
        recognizer.dynamic_energy_threshold = True
        recognizer.pause_threshold = 0.8

        with sr.AudioFile(str(wav_path)) as source:
            print("Reading audio file...")
            # Adjust for ambient noise to improve accuracy
            recognizer.adjust_for_ambient_noise(source, duration=0.5)
            # Record the audio with optimized parameters
            audio = recognizer.record(source)

            print("Sending to speech recognition service...")
            transcript = recognizer.recognize_google(audio, language="en-US")
            print(f"Transcription received. Length: {len(transcript)} characters")
            print(f"Transcript: {transcript[:200]}..." if len(transcript) > 200 else f"Transcript: {transcript}")
            return transcript.lower()
            return transcript.lower()
    except sr.UnknownValueError:
        print(f"Could not understand audio in {audio_path}")
        return ""
    except sr.RequestError as e:
        print(f"Error with speech recognition service for {audio_path}; {e}")
        return ""
    except Exception as e:
        print(f"Unexpected error processing audio: {e}")
        return ""
    finally:
        # Clean up temporary WAV file if it was converted
        try:
            if 'wav_path' in locals() and wav_path != str(audio_path) and Path(wav_path).exists():
                Path(wav_path).unlink(missing_ok=True)
                print("Cleaned up temporary WAV file")
        except Exception as e:
            print(f"Error cleaning up temporary file: {e}")


def format_text(text):
    """
    Apply advanced formatting rules to make transcribed text more professional.
    Uses heuristics to detect sentence boundaries and format accordingly.

    Args:
        text: Raw text string from transcription

    Returns:
        Formatted text with proper capitalization and punctuation
    """
    if not text:
        return text

    # Preprocessing - clean up common transcription artifacts
    text = text.strip()

    # Replace multiple spaces with a single space
    text = ' '.join(text.split())

    # Sentence boundary detection
    # Add periods where they're likely missing between sentences

    # Common indicators of sentence boundaries in speech
    boundary_indicators = [
        ' and then ', ' afterwards ', ' after that ', ' next ', ' following that ',
        ' subsequently ', ' consequently ', ' as a result ', ' therefore ',
        ' thus ', ' hence ', ' so ', ' finally ', ' lastly ', ' in conclusion ',
        ' to sum up ', ' in summary ', ' additionally ', ' moreover ', ' furthermore ',
        ' in addition ', ' also ', ' besides ', ' however ', ' nevertheless ',
        ' nonetheless ', ' still ', ' yet ', ' conversely ', ' on the other hand ',
        ' on the contrary ', ' in contrast ', ' instead ', ' alternatively ',
        ' otherwise ', ' rather ', ' whereas ', ' while ', ' though ', ' although '
    ]

    # Add periods before these indicators if there's no punctuation
    for indicator in boundary_indicators:
        text = text.replace(indicator, f'. {indicator.strip()}')

    # Replace common conjunctions at the start of a sentence that might indicate a boundary
    conjunctions = [' but ', ' or ', ' nor ', ' for ', ' so ']
    for conj in conjunctions:
        text = text.replace(conj, f'. {conj.strip()} ')

    # Split text into sentences based on existing punctuation and the indicators we added
    # This regex will split on ., !, ? followed by a space or end of string
    import re
    sentences = re.split(r'([.!?])\s+', text)

    # Rejoin sentences with proper spacing and formatting
    formatted_text = ''
    i = 0
    while i < len(sentences):
        if i < len(sentences) - 1 and sentences[i + 1] in '.!?':
            # This is a sentence + its ending punctuation
            sentence = sentences[i]
            punct = sentences[i + 1]

            # Capitalize first letter of sentence if it's not already
            if sentence and sentence[0].islower():
                sentence = sentence[0].upper() + sentence[1:]

            formatted_text += sentence + punct + ' '
            i += 2
        else:
            # This might be the last sentence or one without punctuation
            sentence = sentences[i]

            # Capitalize first letter of sentence if it's not already
            if sentence and sentence[0].islower():
                sentence = sentence[0].upper() + sentence[1:]

            # Add period if there's no ending punctuation
            if sentence and not sentence.endswith(('.', '!', '?')):
                sentence += '.'

            formatted_text += sentence + ' '
            i += 1

    # Final cleanup and formatting
    formatted_text = formatted_text.strip()

    # Fix common capitalization issues
    formatted_text = re.sub(r'\bi\b', 'I', formatted_text)  # Capitalize standalone 'i'
    formatted_text = re.sub(r'\bi\'m\b', 'I\'m', formatted_text, flags=re.IGNORECASE)
    formatted_text = re.sub(r'\bi\'ll\b', 'I\'ll', formatted_text, flags=re.IGNORECASE)
    formatted_text = re.sub(r'\bi\'ve\b', 'I\'ve', formatted_text, flags=re.IGNORECASE)
    formatted_text = re.sub(r'\bi\'d\b', 'I\'d', formatted_text, flags=re.IGNORECASE)

    # Capitalize proper nouns (common ones in your context)
    proper_nouns = ['monday', 'tuesday', 'wednesday', 'thursday', 'friday', 'saturday', 'sunday',
                    'january', 'february', 'march', 'april', 'may', 'june', 'july', 'august',
                    'september', 'october', 'november', 'december']

    for noun in proper_nouns:
        formatted_text = re.sub(r'\b' + noun + r'\b', noun.capitalize(),
                                formatted_text, flags=re.IGNORECASE)

    # Fix spacing around punctuation
    formatted_text = re.sub(r'\s+([.,;:!?])', r'\1', formatted_text)

    # Ensure consistent spacing after punctuation
    formatted_text = re.sub(r'([.,;:!?])(\S)', r'\1 \2', formatted_text)

    # Remove double periods
    formatted_text = formatted_text.replace('..', '.')

    # Ensure the text ends with proper punctuation
    if not formatted_text[-1] in '.!?':
        formatted_text += '.'

    return formatted_text

def debug_raw_transcript(audio_path):
    """Debug function to show raw transcription before any processing."""
    transcript = transcribe_audio_optimized(audio_path)
    print("\n--- RAW TRANSCRIPT ---")
    print(transcript)
    print("--- END RAW TRANSCRIPT ---\n")
    return transcript

def correct_transcript_with_gpt(transcript):
    """
    Use GPT-3.5-Turbo to correct property inspection transcription errors
    while strictly preserving the format.

    Args:
        transcript: The raw transcription text from basic speech recognition

    Returns:
        Corrected transcript text with domain-specific terms fixed
    """
    try:
        # Skip processing if transcript is empty
        if not transcript:
            return transcript

        # Log the original transcript for debugging
        logging.info(f"Original transcript: {transcript}")

        # Create a system prompt with property inspection domain knowledge
        # with strict instructions about format preservation
        system_prompt = """You are an expert assistant for property inspectors. 
        Your task is to correct property inspection transcription errors while precisely maintaining the exact format.

        EXTREMELY IMPORTANT: You must preserve all format markers EXACTLY as they appear:
        - "new room" must remain exactly as "new room"
        - "new attribute" must remain exactly as "new attribute"
        - "new feature" must remain exactly as "new feature" 
        - "comment" must remain exactly as "comment"

        These markers control how the transcript is processed, and any change to them will break the processing.

        Only correct real estate terminology, proper nouns, and obvious speech recognition errors:
        - "Germany" should be corrected to "generally" when referring to condition
        - "Mark" to "mark" when not referring to a person's name
        - "tennant" to "tenant"
        - "lanlord" to "landlord"

        DO NOT add punctuation, capitalization, or restructure the text in any way.
        DO NOT combine or split sections marked by these key phrases.
        DO NOT add or remove any "new room", "new attribute", "new feature", or "comment" markers.

        Return ONLY the corrected transcript with minimal changes.
        """

        # Send the transcript to GPT-3.5-Turbo for correction
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user",
                 "content": f"Correct ONLY the terminology errors in this property inspection transcript while preserving the exact format:\n\n{transcript}"}
            ],
            temperature=0.1,  # Very low temperature for more deterministic output
            max_tokens=2000
        )

        # Extract the corrected text
        corrected_transcript = response.choices[0].message.content.strip()

        # Log the corrected transcript for debugging
        logging.info(f"Corrected transcript: {corrected_transcript}")

        # Check if key markers are preserved
        key_markers = ["new room", "new attribute", "new feature", "comment"]
        for marker in key_markers:
            original_count = transcript.lower().count(marker)
            corrected_count = corrected_transcript.lower().count(marker)

            if original_count != corrected_count:
                logging.warning(
                    f"Format marker '{marker}' count changed! Original: {original_count}, Corrected: {corrected_count}")
                return transcript  # Return original if format changed

        return corrected_transcript

    except Exception as e:
        logging.error(f"Error in GPT correction: {str(e)}")
        # Fall back to original transcript if API call fails
        return transcript


def correct_csv_with_gpt(csv_path):
    """
    Process a completed transcript CSV with GPT to correct terminology and speech errors.

    Args:
        csv_path: Path to the CSV file containing transcription data

    Returns:
        Path to the corrected CSV file
    """
    try:
        # Read the CSV file
        df = pd.read_csv(csv_path)

        # Check if we have comments to process
        if 'Comment' not in df.columns:
            logging.warning(f"No 'Comment' column found in {csv_path}")
            return csv_path

        # Process each comment with GPT-4
        logging.info(f"Processing {len(df)} comments with GPT-4")
        for i, row in df.iterrows():
            if pd.notna(row['Comment']) and row['Comment'].strip():
                original_comment = row['Comment']
                corrected_comment = correct_transcript_with_gpt(original_comment)

                # Update the dataframe with corrected comment
                df.at[i, 'Comment'] = corrected_comment

                # Log if there was a change
                if original_comment != corrected_comment:
                    logging.info(f"Correction made: '{original_comment}' → '{corrected_comment}'")

        # Create path for corrected file
        original_filename = Path(csv_path).name
        corrected_filename = f"corrected_{original_filename}"
        corrected_path = Path(csv_path).parent / corrected_filename

        # Save the corrected CSV
        df.to_csv(corrected_path, index=False)
        logging.info(f"Corrected CSV saved to {corrected_path}")

        return corrected_path

    except Exception as e:
        logging.error(f"Error processing CSV with GPT: {str(e)}")
        # Return original path if processing fails
        return csv_path


# Alternative approach: Post-process existing CSVs
def post_process_existing_csv(csv_path):
    """Apply GPT correction to an existing CSV file."""
    return correct_csv_with_gpt(csv_path)


def process_transcript(transcript):
    """
    Process transcript text following exact hierarchical rules:
    * New Room → new row with room, first attribute, first feature, first comment
    * New Attribute → new row with attribute, first feature, first comment
    * New Feature → new row with feature, first comment
    * New Comment → new row with just the comment
    """
    print("\nProcessing transcript...")
    results = []

    # Split transcript into lines for cleaner processing
    lines = transcript.strip().split('\n')
    lines = [line.strip() for line in lines if line.strip()]

    current_room = ""
    current_attribute = ""
    current_feature = ""
    pending_row = {
        "Room": "",
        "Attribute": "",
        "Feature": "",
        "Comment": "",
        "Tenant Responsibility (TR)": ""
    }

    # Process each line
    for line in lines:
        line = line.strip()

        # Skip empty lines
        if not line:
            continue

        lower_line = line.lower()

        # Handle "new room"
        if lower_line.startswith("new room"):
            # Add any pending row first
            if any(pending_row.values()):
                results.append(pending_row.copy())

            # Start a new row with this room
            current_room = line[8:].strip()  # Remove "new room" prefix
            pending_row = {
                "Room": current_room,
                "Attribute": "",
                "Feature": "",
                "Comment": "",
                "Tenant Responsibility (TR)": ""
            }

        # Handle "new attribute"
        elif lower_line.startswith("new attribute"):
            # If we have a pending room but no attribute yet, add to same row
            if pending_row["Room"] and not pending_row["Attribute"]:
                current_attribute = line[13:].strip()  # Remove "new attribute" prefix
                pending_row["Attribute"] = current_attribute
            else:
                # Otherwise, save pending row and start new row
                if any(pending_row.values()):
                    results.append(pending_row.copy())

                current_attribute = line[13:].strip()  # Remove "new attribute" prefix
                pending_row = {
                    "Room": "",  # Empty for hierarchical display
                    "Attribute": current_attribute,
                    "Feature": "",
                    "Comment": "",
                    "Tenant Responsibility (TR)": ""
                }

        # Handle "new feature"
        elif lower_line.startswith("new feature"):
            # If we have a pending attribute but no feature yet, add to same row
            if (pending_row["Room"] or pending_row["Attribute"]) and not pending_row["Feature"]:
                current_feature = line[11:].strip()  # Remove "new feature" prefix
                pending_row["Feature"] = current_feature
            else:
                # Otherwise, save pending row and start new row
                if any(pending_row.values()):
                    results.append(pending_row.copy())

                current_feature = line[11:].strip()  # Remove "new feature" prefix
                pending_row = {
                    "Room": "",  # Empty for hierarchical display
                    "Attribute": "",  # Empty for hierarchical display
                    "Feature": current_feature,
                    "Comment": "",
                    "Tenant Responsibility (TR)": ""
                }

        # Handle "comment"
        elif lower_line.startswith("comment"):
            comment_text = line[7:].strip()  # Remove "comment" prefix
            is_tenant_responsibility = "tenant responsibility" in comment_text.lower()

            # Apply text formatting if function is available
            if callable(format_text):
                try:
                    comment_text = format_text(comment_text)
                except Exception as e:
                    print(f"Warning: Error formatting comment: {str(e)}")

            # If we have a pending feature but no comment yet, add to same row
            if (pending_row["Room"] or pending_row["Attribute"] or pending_row["Feature"]) and not pending_row[
                "Comment"]:
                pending_row["Comment"] = comment_text
                pending_row["Tenant Responsibility (TR)"] = "✓" if is_tenant_responsibility else ""
            else:
                # Otherwise, save pending row and start new row with just the comment
                if any(pending_row.values()):
                    results.append(pending_row.copy())

                pending_row = {
                    "Room": "",  # Empty for hierarchical display
                    "Attribute": "",  # Empty for hierarchical display
                    "Feature": "",  # Empty for hierarchical display
                    "Comment": comment_text,
                    "Tenant Responsibility (TR)": "✓" if is_tenant_responsibility else ""
                }

    # Add final pending row if any
    if any(pending_row.values()):
        results.append(pending_row.copy())

    print(f"Processed transcript into {len(results)} rows")

    return results


import csv

def process_audio_file(file_path, original_filename):
    """Process audio file with enhanced GPT correction."""
    # Get original transcription
    transcript = transcribe_audio_optimized(file_path)

    if not transcript:
        return None, "Failed to transcribe audio"

    # Apply GPT correction
    corrected_transcript = correct_transcript_with_gpt(transcript)

    # Process the corrected transcript
    results = process_transcript(corrected_transcript)

    if not results:
        return None, "No features found in the transcript"

    # Create DataFrame
    df = pd.DataFrame(results)

    # Generate output filename
    base_filename = Path(original_filename).stem
    output_filename = f"{base_filename}_transcript.csv"
    output_path = TRANSCRIPT_FOLDER / output_filename

    # Save to CSV - without using csv.QUOTE_ALL to avoid the error
    # Just use the default pandas settings
    df.to_csv(output_path, index=False)

    return output_path, None

# If you want to keep the quoting behavior but don't want to add the import,
# you can use this alternative:
def process_audio_file_alt(file_path, original_filename):
    """Process audio file with enhanced GPT correction."""
    # Get original transcription
    transcript = transcribe_audio_optimized(file_path)

    if not transcript:
        return None, "Failed to transcribe audio"

    # Apply GPT correction
    corrected_transcript = correct_transcript_with_gpt(transcript)

    # Process the corrected transcript
    results = process_transcript(corrected_transcript)

    if not results:
        return None, "No features found in the transcript"

    # Create DataFrame
    df = pd.DataFrame(results)

    # Generate output filename
    base_filename = Path(original_filename).stem
    output_filename = f"{base_filename}_transcript.csv"
    output_path = TRANSCRIPT_FOLDER / output_filename

    # Save to CSV - use pandas built-in quoting option instead of csv.QUOTE_ALL
    df.to_csv(output_path, index=False, quoting=1)  # 1 is equivalent to csv.QUOTE_ALL

    return output_path, None


# Add a basic test route
@application.route('/test')
def test():
    return "Server is working!"


@application.route('/basic')
def basic():
    return """
    <!DOCTYPE html>
    <html>
    <head>
        <title>PhillyScript Test</title>
        <style>
            body { font-family: sans-serif; padding: 20px; }
            h1 { color: #4a6fa5; }
        </style>
    </head>
    <body>
        <h1>PhillyScript - Test Page</h1>
        <p>If you can see this page, your Flask server is working correctly.</p>
    </body>
    </html>
    """

@application.route('/')
def index():
    try:
        # Serve the index.html page as the root
        return render_template('index.html')
    except Exception as e:
        return f"Error: {str(e)}"

# Add a separate route for the transcription page
@application.route('/transcribe')
@login_required
def transcribe():
    try:
        return render_template('transcription_page.html')
    except Exception as e:
        return f"Error: {str(e)}"


@application.route('/upload', methods=['POST'])
def upload_file():
    if 'audioFile' not in request.files:
        return jsonify({'status': 'error', 'message': 'No file part'})

    file = request.files['audioFile']

    if file.filename == '':
        return jsonify({'status': 'error', 'message': 'No selected file'})

    # Check if file has an allowed extension
    allowed_extensions = ['.mp3', '.wav', '.flac', '.m4a', '.aac']
    file_ext = Path(file.filename).suffix.lower()

    if file_ext not in allowed_extensions:
        return jsonify({
            'status': 'error',
            'message': 'Unsupported file type. Please upload an MP3, WAV, FLAC, M4A, or AAC file.'
        })

    # Generate a unique ID for this upload
    process_id = str(uuid.uuid4())

    # Create temp file path
    temp_file_path = UPLOAD_FOLDER / f"{process_id}_{file.filename}"

    # Save the file
    file.save(temp_file_path)

    # Return process ID so client can check progress
    return jsonify({
        'status': 'success',
        'processId': process_id,
        'message': 'File uploaded successfully. Processing started.'
    })


@application.route('/process/<process_id>', methods=['POST'])
def process_file(process_id):
    # Find the uploaded file
    uploaded_files = list(UPLOAD_FOLDER.glob(f"{process_id}_*"))

    if not uploaded_files:
        return jsonify({'status': 'error', 'message': 'No file found for this process ID'})

    file_path = uploaded_files[0]
    original_filename = file_path.name.replace(f"{process_id}_", "")

    try:
        # First optimize the audio file
        optimized_path, error_msg = process_uploaded_audio(str(file_path), original_filename)

        if error_msg:
            return jsonify({'status': 'error', 'message': error_msg})

        # Now process with the optimized file
        result_path, error_msg = process_audio_file(str(optimized_path), original_filename)

        if error_msg:
            return jsonify({'status': 'error', 'message': error_msg})

        # Return success with the output filename
        return jsonify({
            'status': 'success',
            'outputFilename': Path(result_path).name,
            'message': 'Processing complete'
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})
    finally:
        # Clean up uploaded file
        try:
            file_path.unlink(missing_ok=True)
        except:
            pass


@application.route('/correct_csv/<filename>', methods=['POST'])
@login_required
def correct_csv(filename):
    """Route to apply GPT correction to an existing CSV file."""
    try:
        # Find the CSV file
        csv_path = TRANSCRIPT_FOLDER / filename
        if not csv_path.exists():
            return jsonify({
                'status': 'error',
                'message': 'CSV file not found'
            })

        # Apply GPT correction
        corrected_path = post_process_existing_csv(csv_path)

        # Return path to corrected file
        return jsonify({
            'status': 'success',
            'outputFilename': Path(corrected_path).name,
            'message': 'CSV corrected successfully'
        })
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': f'Error correcting CSV: {str(e)}'
        })

@application.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    file_path = TRANSCRIPT_FOLDER / filename
    if not file_path.exists():
        return jsonify({'status': 'error', 'message': 'File not found'})

    # Return the file for download
    return send_file(file_path, as_attachment=True)


@application.route('/diff_check')
@login_required
def diff_check():
    """Render the image difference checker page"""
    return render_template('diff_check.html')


@application.route('/diff_result/<result_id>')
def diff_result(result_id):
    """Render the results page for image comparison"""
    return render_template('diff_result.html')


@application.route('/api/check_files/<result_id>')
def check_files(result_id):
    """Check if result files exist"""
    original_path = RESULT_FOLDER / f"{result_id}_original.jpg"
    diff_path = RESULT_FOLDER / f"{result_id}_diff.jpg"

    return jsonify({
        'result_folder': str(RESULT_FOLDER),
        'original_exists': original_path.exists(),
        'diff_exists': diff_path.exists(),
        'files_in_folder': [f.name for f in RESULT_FOLDER.iterdir() if f.is_file()][:10]  # List first 10 files
    })
@application.route('/results/<filename>')
def serve_result_file(filename):
    """Serve files from the results folder"""
    return send_from_directory(RESULT_FOLDER, filename)

@application.route('/api/diff_result/<result_id>')
def get_diff_result(result_id):
    """API endpoint to get the image comparison results"""
    # Check if result exists
    original_path = RESULT_FOLDER / f"{result_id}_original.jpg"
    diff_path = RESULT_FOLDER / f"{result_id}_diff.jpg"

    if not original_path.exists() or not diff_path.exists():
        return jsonify({
            'status': 'error',
            'message': 'Result not found'
        })

    # Read the metadata file for difference count
    metadata_path = RESULT_FOLDER / f"{result_id}_metadata.txt"
    difference_count = 0
    if metadata_path.exists():
        with open(metadata_path, 'r') as f:
            metadata = f.read().strip()
            try:
                difference_count = int(metadata)
            except:
                difference_count = 0

    # Return paths and metadata with UPDATED URL PATHS
    return jsonify({
        'status': 'success',
        'originalImageUrl': f"/results/{result_id}_original.jpg",  # Updated path
        'diffImageUrl': f"/results/{result_id}_diff.jpg",         # Updated path
        'differenceCount': difference_count
    })


@application.route('/compare_images', methods=['POST'])
def compare_images():
    """Endpoint to handle image comparison uploads and processing"""
    if 'image1' not in request.files or 'image2' not in request.files:
        return jsonify({
            'status': 'error',
            'message': 'Both images are required'
        })

    file1 = request.files['image1']
    file2 = request.files['image2']

    if file1.filename == '' or file2.filename == '':
        return jsonify({
            'status': 'error',
            'message': 'No file selected'
        })

    # Generate a unique ID for this comparison
    result_id = str(uuid.uuid4())

    # Save uploaded files temporarily
    temp_path1 = UPLOAD_FOLDER / f"{result_id}_1{Path(secure_filename(file1.filename)).suffix}"
    temp_path2 = UPLOAD_FOLDER / f"{result_id}_2{Path(secure_filename(file2.filename)).suffix}"

    file1.save(temp_path1)
    file2.save(temp_path2)

    # Process images and find differences
    try:
        difference_count = process_image_difference(
            str(temp_path1),
            str(temp_path2),
            str(RESULT_FOLDER / f"{result_id}_original.jpg"),
            str(RESULT_FOLDER / f"{result_id}_diff.jpg")
        )

        # Save metadata
        with open(RESULT_FOLDER / f"{result_id}_metadata.txt", 'w') as f:
            f.write(str(difference_count))

        # Clean up temporary files
        temp_path1.unlink(missing_ok=True)
        temp_path2.unlink(missing_ok=True)

        return jsonify({
            'status': 'success',
            'resultId': result_id,
            'message': 'Images compared successfully'
        })

    except Exception as e:
        # Clean up temporary files
        temp_path1.unlink(missing_ok=True)
        temp_path2.unlink(missing_ok=True)

        return jsonify({
            'status': 'error',
            'message': f'Error processing images: {str(e)}'
        })


def process_image_difference(image1_path, image2_path, output_original_path, output_diff_path):
    """
    Process two images to find and highlight differences

    Args:
        image1_path: Path to the reference image
        image2_path: Path to the comparison image
        output_original_path: Path where to save the original image
        output_diff_path: Path where to save the diff image with highlighted differences

    Returns:
        The number of differences found
    """
    # Load images
    img1 = cv2.imread(image1_path)
    img2 = cv2.imread(image2_path)

    # Check if images were loaded correctly
    if img1 is None or img2 is None:
        raise ValueError("Failed to load one or both images")

    # Resize images to match if they have different dimensions
    if img1.shape != img2.shape:
        # Resize the second image to match the first
        img2 = cv2.resize(img2, (img1.shape[1], img1.shape[0]))

    # Convert images to grayscale
    gray1 = cv2.cvtColor(img1, cv2.COLOR_BGR2GRAY)
    gray2 = cv2.cvtColor(img2, cv2.COLOR_BGR2GRAY)

    # Compute absolute difference between grayscale images
    diff = cv2.absdiff(gray1, gray2)

    # Apply threshold to highlight differences
    _, thresh = cv2.threshold(diff, 30, 255, cv2.THRESH_BINARY)

    # Find contours of differences
    contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    # Filter contours by size (remove very small differences that might be noise)
    min_area = 50  # Minimum area in pixels
    significant_contours = [cnt for cnt in contours if cv2.contourArea(cnt) > min_area]

    # Create a copy of the second image to draw the differences on
    img_diff = img2.copy()

    # Draw rectangles around detected differences
    for contour in significant_contours:
        x, y, w, h = cv2.boundingRect(contour)
        cv2.rectangle(img_diff, (x, y), (x + w, y + h), (0, 255, 0), 2)

    # Save the images
    cv2.imwrite(output_original_path, img1)
    cv2.imwrite(output_diff_path, img_diff)

    return len(significant_contours)


@application.route('/finalise_report')
@login_required
def finalise_report():
    """Render the finalise report page"""
    return render_template('finalise_report.html')


@application.route('/compare_text', methods=['POST'])
@login_required
def compare_text():
    """Endpoint to handle text comparison uploads and processing"""
    if 'original' not in request.files or 'comparison' not in request.files:
        return jsonify({
            'status': 'error',
            'message': 'Both files are required'
        })

    original_file = request.files['original']
    comparison_file = request.files['comparison']

    if original_file.filename == '' or comparison_file.filename == '':
        return jsonify({
            'status': 'error',
            'message': 'No file selected'
        })

    # Generate a unique ID for this comparison
    result_id = str(uuid.uuid4())

    # Save uploaded files temporarily
    original_path = UPLOAD_FOLDER / f"{result_id}_original{Path(secure_filename(original_file.filename)).suffix}"
    comparison_path = UPLOAD_FOLDER / f"{result_id}_comparison{Path(secure_filename(comparison_file.filename)).suffix}"

    original_file.save(original_path)
    comparison_file.save(comparison_path)

    # Process files and find differences
    try:
        # Extract text from files
        original_text = extract_text_from_file(original_path)
        comparison_text = extract_text_from_file(comparison_path)

        if not original_text or not comparison_text:
            raise ValueError("Could not extract text from one or both files")

        # Compare texts and generate HTML with differences highlighted
        result_html = generate_diff_html(original_text, comparison_text)

        # Save the result
        result_path = RESULT_FOLDER / f"{result_id}_diff.html"
        with open(result_path, 'w', encoding='utf-8') as f:
            f.write(result_html)

        # Clean up temporary files
        original_path.unlink(missing_ok=True)
        comparison_path.unlink(missing_ok=True)

        return jsonify({
            'status': 'success',
            'resultUrl': f"/view_comparison/{result_id}",  # Changed URL to point to a new route
            'message': 'Files compared successfully'
        })

    except Exception as e:
        # Clean up temporary files
        original_path.unlink(missing_ok=True)
        comparison_path.unlink(missing_ok=True)

        return jsonify({
            'status': 'error',
            'message': f'Error processing files: {str(e)}'
        })


@application.route('/view_comparison/<result_id>')
def view_comparison(result_id):
    """Serve the comparison result directly"""
    result_path = RESULT_FOLDER / f"{result_id}_diff.html"

    if not result_path.exists():
        return "Comparison result not found", 404

    with open(result_path, 'r', encoding='utf-8') as f:
        html_content = f.read()

    return html_content


@application.route('/generate_report', methods=['POST'])
@login_required
def generate_report():
    """Endpoint to handle report generation from CSV"""
    use_latest = request.form.get('useLatest') == 'true'
    report_type = request.form.get('reportType', 'full')  # Default to full if not specified

    try:
        csv_path = None

        if use_latest:
            # Find the most recent CSV file in the transcripts folder
            transcript_dir = Path('/tmp/temp_transcripts')
            if not transcript_dir.exists():
                return jsonify({
                    'status': 'error',
                    'message': 'No transcript directory found'
                })

            csv_files = list(transcript_dir.glob('*.csv'))
            if not csv_files:
                return jsonify({
                    'status': 'error',
                    'message': 'No CSV files found in the transcript directory'
                })

            # Sort by modification time, newest first
            csv_files.sort(key=lambda x: x.stat().st_mtime, reverse=True)
            csv_path = csv_files[0]
        else:
            if 'csv' not in request.files:
                return jsonify({
                    'status': 'error',
                    'message': 'No CSV file provided'
                })

            csv_file = request.files['csv']
            if csv_file.filename == '':
                return jsonify({
                    'status': 'error',
                    'message': 'No file selected'
                })

            # Generate a unique ID for this report
            result_id = str(uuid.uuid4())

            # Save uploaded file temporarily
            csv_path = UPLOAD_FOLDER / f"{result_id}_{secure_filename(csv_file.filename)}"
            csv_file.save(csv_path)

        # Generate report from CSV with the specified report type
        # For now, all report types use the same generation function
        # In the future, you can implement different generation logic based on report_type

        # Log the report type being generated
        logger.info(f"Generating report of type: {report_type}")

        result_path = generate_enhanced_docx_report(csv_path, report_type)

        # Clean up temporary file if it was uploaded
        if not use_latest:
            csv_path.unlink(missing_ok=True)

        return jsonify({
            'status': 'success',
            'resultUrl': f"/download_report/{result_path.name}",
            'message': f'Report generated successfully'
        })

    except Exception as e:
        # Clean up temporary file if it exists and was uploaded
        if csv_path and not use_latest and csv_path.exists():
            csv_path.unlink(missing_ok=True)

        logger.error(f"Error generating report: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': f'Error generating report: {str(e)}'
        })


def generate_report_boilerplate(doc, report_type, address, inspection_date, on_behalf_of):
    """
    Generate standardized boilerplate content for different report types.

    Args:
        doc: The Document object to add content to
        report_type: Type of report ('inventory', 'full', or 'checkout')
        address: Property address
        inspection_date: Date of inspection (formatted string)
        on_behalf_of: Preparer name
    """
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Add logo (placeholder)
    # In a real implementation, you would add a company logo here
    logo_paragraph = doc.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_run = logo_paragraph.add_run("Ft² Inventories")
    logo_run.bold = True
    logo_run.font.size = Pt(16)

    # Add address on the right side
    address_paragraph = doc.add_paragraph()
    address_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    address_run = address_paragraph.add_run(address)

    # Add main report titles based on report type
    if report_type == 'inventory':
        # INVENTORY CHECK-IN REPORT FORMAT
        title = doc.add_heading('INVENTORY', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add And
        and_heading = doc.add_heading('And', 0)
        and_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add CHECK IN
        checkin_heading = doc.add_heading('CHECK IN', 0)
        checkin_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add subtitle
        subtitle = doc.add_paragraph('Of the contents and conditions for')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add property address in a box
        address_box_para = doc.add_paragraph()
        address_box_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        address_box = address_box_para.add_run(address)

        # Add date of inspection
        date_para = doc.add_paragraph(f"Date of inspection: {inspection_date}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add prepared by
        prepared_para = doc.add_paragraph(f"PREPARED BY: Ft² Inventories Ltd")
        prepared_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add company address
        company_address = doc.add_paragraph("Birch Tree House, Glympton Road, Wootton, Woodstock, OX20 1EJ")
        company_address.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add company contact
        company_contact = doc.add_paragraph("info@ft2inventories.co.uk  020 8004 3324")
        company_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add on behalf of section
        behalf_para = doc.add_paragraph("On behalf of:")
        behalf_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add name in a table
        behalf_table = doc.add_table(rows=1, cols=2)
        behalf_table.autofit = False
        behalf_table.columns[0].width = Inches(1)
        behalf_table.columns[1].width = Inches(5)
        behalf_cells = behalf_table.rows[0].cells
        behalf_cells[0].text = "Name"
        behalf_cells[1].text = on_behalf_of

        # Add important information heading
        info_heading = doc.add_heading('Important Information', level=1)
        info_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add important information text
        info_text = doc.add_paragraph(
            "The Tenant /Tenant's representative and landlord should sign this document to signify that they have read and understood each page and accept that this Inventory and Check in is a true and accurate representation at the date so specified of the Property at the address stated above. Amendments to this document may be made within 5 days from the date the report is sent out. If no amendments are made within this time the parties will be deemed to accept the document as an accurate representation at the date so specified of the property at the address so stated above without the need for signature. This document should be returned, signed and dated no later than 5 days from the date the report is sent out.")

        # Add table of contents heading
        toc_heading = doc.add_heading('Table of Contents', level=1)

        # Add placeholder table of contents
        toc = doc.add_table(rows=6, cols=2)
        toc.style = 'Table Grid'

        # Add table of contents rows
        toc_rows = [
            ("Contents", "Page number"),
            ("Disclaimer and important information", "[tc1]"),
            ("General description of conditions/ Utility readings", "[tc2]"),
            ("Front Door and Entrance Hall", "[tc3]"),
            ("Bedroom", "[tc4]"),
            ("Bathroom", "[tc5]")
        ]

        for i, (content, page) in enumerate(toc_rows):
            cells = toc.rows[i].cells
            cells[0].text = content
            cells[1].text = page

        # Add a page break before disclaimer section
        doc.add_page_break()

        # Add INVENTORY DISCLAIMERS heading
        disclaimer_heading = doc.add_heading('INVENTORY DISCLAIMERS', level=1)
        disclaimer_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add disclaimer table
        disclaimer_table = doc.add_table(rows=10, cols=2)
        disclaimer_table.style = 'Table Grid'

        # Add disclaimer content (just a few examples)
        disclaimer_rows = [
            ("1.", "Structural", "This Inventory does not constitute a structural survey of the Property."),
            ("2.", "General",
             "This Inventory has been prepared on the accepted principle that all items are free from any obvious damage, fault or soiling except where stated. The term 'good' is noted as a guideline for this. NOTE: Should there be any cleaning, health and safety or other issues that need urgent attention, please call the Agent's office immediately."),
            ("3.", "Description",
             "Where the words 'gold', 'brass', 'oak', 'walnut' etc are used, it is understood that this is a description of the colour and type of the item and not the actual fabric, unless documentary evidence is available."),
            ("4.", "Attendees",
             "The Clerk must be alone in the property for the inventory inspection. The tenant / Landlord or a representative is welcome to briefly attend the inspection should they wish to do so."),
            ("5.", "Fire Safety Equipment",
             "If smoke detectors/carbon monoxide monitors are present and replacement batteries are required between maintenance visits or periodic tenancy checks, it is the Tenant's responsibility to replace and frequently check the working order of the same.")
        ]

        for i, (num, title, desc) in enumerate(disclaimer_rows[:5]):
            cells = disclaimer_table.rows[i].cells
            cells[0].text = title
            cells[1].text = desc

    elif report_type == 'full':
        # FULL CHECK-IN REPORT FORMAT
        title = doc.add_heading('Full Check-In Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add property address
        address_para = doc.add_paragraph(address)
        address_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add date
        date_paragraph = doc.add_paragraph()
        date_run = date_paragraph.add_run(f"Inspection Date: {inspection_date}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add prepared by
        prepared_para = doc.add_paragraph(f"Prepared by: {on_behalf_of}")
        prepared_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add introduction
        doc.add_heading('Introduction', level=1)
        doc.add_paragraph(
            "This full check-in report documents the complete condition of the property at the beginning of the tenancy and identifies any issues that require attention. Items marked with 'TR' indicate tenant responsibility.")

    elif report_type == 'checkout':
        # CHECK-OUT REPORT FORMAT
        title = doc.add_heading('Check-Out Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add property address
        address_para = doc.add_paragraph(address)
        address_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add date
        date_paragraph = doc.add_paragraph()
        date_run = date_paragraph.add_run(f"Inspection Date: {inspection_date}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add prepared by
        prepared_para = doc.add_paragraph(f"Prepared by: {on_behalf_of}")
        prepared_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add introduction
        doc.add_heading('Introduction', level=1)
        doc.add_paragraph(
            "This check-out report documents the condition of the property at the end of the tenancy and identifies any changes from the check-in report. Items marked with 'TR' indicate tenant responsibility.")

    # Add page break after boilerplate
    doc.add_page_break()

    # Add footer with date to every page
    # Note: This would require more complex handling with python-docx
    # For now, we'll just add a placeholder at the bottom of each page
    footer_para = doc.add_paragraph(f"Date of inspection: {inspection_date}")
    footer_para.style = 'Footer'


def generate_report_closing(doc, report_type, page_count):
    """
    Generate standardized closing content for different report types.

    Args:
        doc: The Document object to add content to
        report_type: Type of report ('inventory', 'full', or 'checkout')
        page_count: Total number of pages in the document (for declaration)
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Add a page break before closing sections
    doc.add_page_break()

    # Add closing content based on report type
    if report_type in ['inventory', 'full']:
        # Add tenant and landlord declaration section
        doc.add_heading('TENANT DECLARATION', level=1)
        doc.add_paragraph(
            f"The items listed in all [x-pages] of this inventory/check in have been inspected and found to be in the condition indicated.")
        doc.add_paragraph(
            "I ………………………………………. being of sound mind have fully understood the implications of signing this document and verify that the content is correct and accurate at the time of signing and that the content will be binding if relied upon in a Court of Law.")

        # Add signature fields
        doc.add_paragraph("Name …………………………………………………………")
        doc.add_paragraph("Signed for the Tenant ……………………………………………")

        doc.add_heading('LANDLORD DECLARATION', level=1)
        doc.add_paragraph(
            f"The items listed in all 124 pages of this inventory/check in have been inspected and found to be in the condition indicated.")
        doc.add_paragraph(
            "I ………………………………………. being of sound mind have fully understood the implications of signing this document and verify that the content is correct and accurate at the time of signing and that the content will be binding if relied upon in a Court of Law.")

        # Add signature fields
        doc.add_paragraph("Name ……………………………………………………………")
        doc.add_paragraph("Signed for the Landlord …………………………………………")

    elif report_type == 'checkout':
        # Add summary
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(
            "This check-out report provides a comprehensive overview of the property's condition at the end of the tenancy. Any discrepancies with the check-in report have been noted.")

        # Add signatures for checkout
        doc.add_heading('CHECKOUT CONFIRMATION', level=1)
        doc.add_paragraph(
            "This checkout report has been completed and represents the condition of the property at the end of the tenancy.")

        # Add signature fields
        doc.add_paragraph("Tenant Name: …………………………………………………………")
        doc.add_paragraph("Tenant Signature: ……………………………………………")
        doc.add_paragraph("Date: ……………………………………")

        doc.add_paragraph("Agent/Landlord Name: …………………………………………………………")
        doc.add_paragraph("Agent/Landlord Signature: ……………………………………………")
        doc.add_paragraph("Date: ……………………………………")

@application.route('/download_report/<filename>')
def download_report(filename):
    """Serve the generated report for download"""
    return send_from_directory(RESULT_FOLDER, filename, as_attachment=True)


def extract_text_from_file(file_path):
    """Extract text content from various file formats"""
    file_ext = file_path.suffix.lower()

    try:
        if file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()

        elif file_ext in ['.doc', '.docx']:
            # Use the properly imported Document class
            from docx import Document
            doc = Document(file_path)
            return '\n'.join([para.text for para in doc.paragraphs])

        elif file_ext == '.pdf':
            # Note: This would require PyPDF2 or another PDF library
            # For simplicity, we'll just return an error for now
            raise ValueError("PDF extraction not implemented")

        elif file_ext == '.rtf':
            # RTF parsing is complex, you'd need a library like striprtf
            raise ValueError("RTF extraction not implemented")

        else:
            # Try to read as plain text by default
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()

    except Exception as e:
        raise ValueError(f"Could not extract text from file: {str(e)}")


def generate_diff_html(original_text, comparison_text):
    """
    Generate HTML highlighting differences between two texts
    """
    # Split texts into lines
    original_lines = original_text.splitlines()
    comparison_lines = comparison_text.splitlines()

    # Get differences
    differ = difflib.HtmlDiff(wrapcolumn=80)
    diff_html = differ.make_file(original_lines, comparison_lines,
                                 'Original Text', 'Comparison Text',
                                 context=True, numlines=3)

    # Enhance the HTML with our own styling
    styled_html = f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Text Comparison Results</title>
        <style>
            body {{
                font-family: 'Segoe UI', Arial, sans-serif;
                margin: 20px;
                line-height: 1.5;
            }}
            .diff-header {{
                background-color: #4a6fa5;
                color: white;
                padding: 10px;
                border-radius: 5px;
                margin-bottom: 20px;
                display: flex;
                justify-content: space-between;
                align-items: center;
            }}
            .diff-header h1 {{
                margin: 0;
                font-size: 24px;
            }}
            table.diff {{
                width: 100%;
                border-collapse: collapse;
                font-size: 14px;
                font-family: monospace;
            }}
            .diff td {{
                padding: 5px;
                border: 1px solid #ddd;
                vertical-align: top;
            }}
            .diff span.diff_add {{
                background-color: #e6ffe6;
                color: green;
                font-weight: bold;
            }}
            .diff span.diff_sub {{
                background-color: #ffe6e6;
                color: red;
                text-decoration: line-through;
            }}
            .diff span.diff_chg {{
                background-color: #fff5cc;
                color: #996600;
                font-weight: bold;
            }}
            .diff th {{
                background-color: #f0f0f0;
                padding: 5px;
                border: 1px solid #ddd;
                text-align: center;
            }}
            .legend {{
                margin-top: 20px;
                padding: 10px;
                background-color: #f9f9f9;
                border-radius: 5px;
            }}
            .legend ul {{
                padding-left: 20px;
            }}
            .back-button {{
                padding: 8px 15px;
                background-color: #4a6fa5;
                color: white;
                border: none;
                border-radius: 5px;
                cursor: pointer;
                text-decoration: none;
                display: inline-block;
                font-size: 14px;
            }}
            .back-button:hover {{
                background-color: #375d8a;
            }}
            .footer {{
                margin-top: 30px;
                text-align: center;
            }}
        </style>
    </head>
    <body>
        <div class="diff-header">
            <h1>Text Comparison Results</h1>
            <button onclick="goBack()" class="back-button">Back to Report Options</button>
        </div>

        {diff_html}

        <div class="legend">
            <h3>Legend:</h3>
            <ul>
                <li><span style="color: green; font-weight: bold;">Added content</span> - Content that appears in the comparison text but not in the original</li>
                <li><span style="color: red; text-decoration: line-through;">Removed content</span> - Content that appears in the original text but not in the comparison</li>
                <li><span style="color: #996600; font-weight: bold;">Changed content</span> - Content that has been modified between the texts</li>
            </ul>
        </div>

        <div class="footer">
            <button onclick="goBack()" class="back-button">Back to Report Options</button>
        </div>

        <script>
        function goBack() {{
            window.location.href = window.location.origin + '/finalise_report';
        }}
        </script>
    </body>
    </html>
    """

    return styled_html


# Add these routes to your application.py file

@application.route('/report_builder')
@login_required
def enhanced_report_builder():
    """Render the enhanced report builder page"""
    return render_template('report_builder.html')


@application.route('/api/get_rooms')
@login_required
def get_rooms():
    """API endpoint to get rooms from a CSV file"""
    csv_id = request.args.get('csvId')

    if not csv_id:
        return jsonify({
            'status': 'error',
            'message': 'Missing csvId parameter'
        })

    # Find the CSV file
    csv_path = None
    if csv_id == 'latest':
        # Find the most recent CSV file
        transcript_dir = Path('/tmp/temp_transcripts')
        if not transcript_dir.exists():
            return jsonify({
                'status': 'error',
                'message': 'No transcript directory found'
            })

        csv_files = list(transcript_dir.glob('*.csv'))
        if not csv_files:
            return jsonify({
                'status': 'error',
                'message': 'No CSV files found'
            })

        # Sort by modification time, newest first
        csv_files.sort(key=lambda x: x.stat().st_mtime, reverse=True)
        csv_path = csv_files[0]
    else:
        # Find the specific CSV file
        csv_path = UPLOAD_FOLDER / f"{csv_id}.csv"
        if not csv_path.exists():
            return jsonify({
                'status': 'error',
                'message': 'CSV file not found'
            })

    try:
        # Read the CSV file
        df = pd.read_csv(csv_path)

        # Extract unique room names (non-empty ones)
        rooms = []
        current_room = None

        for _, row in df.iterrows():
            room = row['Room'].strip() if pd.notna(row['Room']) and row['Room'].strip() else None
            if room:
                current_room = room
                if current_room not in rooms:
                    rooms.append(current_room)

        return jsonify({
            'status': 'success',
            'rooms': rooms
        })
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': f'Error processing CSV: {str(e)}'
        })


@application.route('/api/prepare_report', methods=['POST'])
@login_required
def prepare_report():
    """API endpoint to prepare a report and get the intermediate page"""
    report_type = request.form.get('reportType', 'full')
    use_latest = request.form.get('useLatest') == 'true'

    try:
        csv_path = None
        csv_id = None

        if use_latest:
            # Find the most recent CSV file
            transcript_dir = Path('/tmp/temp_transcripts')
            if not transcript_dir.exists():
                return jsonify({
                    'status': 'error',
                    'message': 'No transcript directory found'
                })

            csv_files = list(transcript_dir.glob('*.csv'))
            if not csv_files:
                return jsonify({
                    'status': 'error',
                    'message': 'No CSV files found'
                })

            # Sort by modification time, newest first
            csv_files.sort(key=lambda x: x.stat().st_mtime, reverse=True)
            csv_path = csv_files[0]
            csv_id = 'latest'
        else:
            if 'csv' not in request.files:
                return jsonify({
                    'status': 'error',
                    'message': 'No CSV file provided'
                })

            csv_file = request.files['csv']
            if csv_file.filename == '':
                return jsonify({
                    'status': 'error',
                    'message': 'No file selected'
                })

            # Generate a unique ID for this CSV
            csv_id = str(uuid.uuid4())

            # Save uploaded file
            csv_path = UPLOAD_FOLDER / f"{csv_id}.csv"
            csv_file.save(csv_path)

        # Redirect to the report builder with the CSV ID and report type
        # Fixed URL to match your route name
        return jsonify({
            'status': 'success',
            'redirectUrl': f'/report_builder?csvId={csv_id}&type={report_type}',
            'message': 'CSV processed successfully'
        })

    except Exception as e:
        # Clean up temporary file if it exists and was uploaded
        if csv_path and not use_latest and csv_path.exists():
            csv_path.unlink(missing_ok=True)

        return jsonify({
            'status': 'error',
            'message': f'Error processing CSV: {str(e)}'
        })


@application.route('/api/generate_enhanced_report', methods=['POST'])
@login_required
def generate_enhanced_report():
    """API endpoint to generate a final report with property details and room images"""
    try:
        # Get form data
        report_type = request.form.get('reportType', 'full')
        csv_id = request.form.get('csvId')
        address = request.form.get('address')
        inspection_date = request.form.get('inspectionDate')
        on_behalf_of = request.form.get('onBehalfOf')

        if not csv_id or not address or not inspection_date or not on_behalf_of:
            return jsonify({
                'status': 'error',
                'message': 'Missing required fields'
            })

        # Find the CSV file
        csv_path = None
        if csv_id == 'latest':
            # Find the most recent CSV file
            transcript_dir = Path('/tmp/temp_transcripts')
            if not transcript_dir.exists():
                return jsonify({
                    'status': 'error',
                    'message': 'No transcript directory found'
                })

            csv_files = list(transcript_dir.glob('*.csv'))
            if not csv_files:
                return jsonify({
                    'status': 'error',
                    'message': 'No CSV files found'
                })

            # Sort by modification time, newest first
            csv_files.sort(key=lambda x: x.stat().st_mtime, reverse=True)
            csv_path = csv_files[0]
        else:
            # Find the specific CSV file
            csv_path = UPLOAD_FOLDER / f"{csv_id}.csv"
            if not csv_path.exists():
                return jsonify({
                    'status': 'error',
                    'message': 'CSV file not found'
                })

        # Process uploaded images for each room - UPDATED VERSION
        room_images = {}

        # Debug: print all keys in request.form and request.files
        logging.info(f"Form keys: {list(request.form.keys())}")
        logging.info(f"File keys: {list(request.files.keys())}")
        logging.info(f"Raw request.files: {request.files}")

        # Check for image files in the request with the new naming convention: roomImages_RoomName
        for key in request.files.keys():
            if key.startswith('roomImages_'):
                # Extract room name from the input name format: roomImages_RoomName
                room_name = key[len('roomImages_'):].lower()  # Extract what's after 'roomImages_'
                logging.info(f"Processing images for room: {room_name}")

                # Initialize room in the dictionary if needed
                if room_name not in room_images:
                    room_images[room_name] = []

                # Get all files for this room
                files = request.files.getlist(key)
                logging.info(f"Number of files for {room_name}: {len(files)}")

                # Save each image file
                for file in files:
                    if file and file.filename:
                        # Generate a unique filename
                        img_id = str(uuid.uuid4())
                        img_ext = Path(secure_filename(file.filename)).suffix
                        img_path = UPLOAD_FOLDER / f"{img_id}{img_ext}"

                        # Save the image
                        file.save(img_path)
                        room_images[room_name].append(str(img_path))

                        # Log the save
                        logging.info(f"Saved image for {room_name}: {img_path} (from {file.filename})")

        # Also check for the original format (as a fallback)
        for key in request.files.keys():
            if key.startswith('roomImages['):
                # Extract room name from the input name format: roomImages[Room Name]
                room_name = key[11:-1].lower()  # Extract what's between 'roomImages[' and ']'
                logging.info(f"Processing images for room (bracket format): {room_name}")

                # Initialize room in the dictionary if needed
                if room_name not in room_images:
                    room_images[room_name] = []

                # Get all files for this room
                files = request.files.getlist(key)
                logging.info(f"Number of files for {room_name}: {len(files)}")

                # Save each image file
                for file in files:
                    if file and file.filename:
                        # Generate a unique filename
                        img_id = str(uuid.uuid4())
                        img_ext = Path(secure_filename(file.filename)).suffix
                        img_path = UPLOAD_FOLDER / f"{img_id}{img_ext}"

                        # Save the image
                        file.save(img_path)
                        room_images[room_name].append(str(img_path))

                        # Log the save
                        logging.info(f"Saved image for {room_name}: {img_path} (from {file.filename})")

        # Debug log the final image count per room
        for room, images in room_images.items():
            logging.info(f"Room {room}: {len(images)} images collected")

        # Generate the enhanced report
        result_path = generate_enhanced_docx_report(
            csv_path,
            report_type,
            address,
            inspection_date,
            on_behalf_of,
            room_images
        )

        # Create a filename based on address and date
        formatted_date = datetime.strptime(inspection_date, '%Y-%m-%d').strftime('%d-%m-%Y')
        report_name = f"{address.replace(' ', '_')}_{formatted_date}_{report_type}_report.docx"

        return jsonify({
            'status': 'success',
            'reportUrl': f"/download_report/{result_path.name}",
            'reportName': report_name,
            'message': f'{report_type.capitalize()} report generated successfully'
        })

    except Exception as e:
        logging.error(f"Error generating report: {str(e)}", exc_info=True)
        return jsonify({
            'status': 'error',
            'message': f'Error generating report: {str(e)}'
        })


@application.route('/api/get_csv_rooms')
@login_required
def get_csv_rooms():
    """API endpoint to get rooms from a CSV file"""
    csv_id = request.args.get('csvId')

    if not csv_id:
        return jsonify({
            'status': 'error',
            'message': 'Missing csvId parameter'
        })

    # Find the CSV file
    csv_path = None
    if csv_id == 'latest':
        # Find the most recent CSV file
        transcript_dir = Path('/tmp/temp_transcripts')
        if not transcript_dir.exists():
            return jsonify({
                'status': 'error',
                'message': 'No transcript directory found'
            })

        csv_files = list(transcript_dir.glob('*.csv'))
        if not csv_files:
            return jsonify({
                'status': 'error',
                'message': 'No CSV files found'
            })

        # Sort by modification time, newest first
        csv_files.sort(key=lambda x: x.stat().st_mtime, reverse=True)
        csv_path = csv_files[0]
    else:
        # Find the specific CSV file - try both locations
        csv_path_upload = UPLOAD_FOLDER / f"{csv_id}.csv"
        csv_path_transcript = TRANSCRIPT_FOLDER / f"{csv_id}.csv"

        if csv_path_upload.exists():
            csv_path = csv_path_upload
        elif csv_path_transcript.exists():
            csv_path = csv_path_transcript
        else:
            # Try to see if the full filename was passed
            for folder in [UPLOAD_FOLDER, TRANSCRIPT_FOLDER]:
                potential_path = folder / csv_id
                if potential_path.exists():
                    csv_path = potential_path
                    break

            if not csv_path:
                return jsonify({
                    'status': 'error',
                    'message': 'CSV file not found'
                })

    try:
        # Read the CSV file
        df = pd.read_csv(csv_path)

        # Extract unique room names (non-empty ones)
        rooms = []
        current_room = None

        for _, row in df.iterrows():
            room_col = 'Room' if 'Room' in df.columns else None
            if not room_col:
                return jsonify({
                    'status': 'error',
                    'message': 'Room column not found in CSV'
                })

            room = row[room_col].strip() if pd.notna(row[room_col]) and row[room_col].strip() else None
            if room:
                current_room = room
                if current_room not in rooms:
                    rooms.append(current_room)

        # Convert room names to lowercase for consistency
        rooms = [room.lower() for room in rooms]

        return jsonify({
            'status': 'success',
            'rooms': rooms
        })
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': f'Error processing CSV: {str(e)}'
        })

def generate_enhanced_docx_report(csv_path, report_type, address, inspection_date, on_behalf_of, room_images):
    """
    Generate an enhanced report with property details and room images

    Args:
        csv_path: Path to the CSV file
        report_type: Type of report to generate ('inventory', 'full', or 'checkout')
        address: Property address
        inspection_date: Date of inspection
        on_behalf_of: Prepared on behalf of
        room_images: Dictionary mapping room names to lists of image paths

    Returns:
        Path to the generated report
    """
    import pandas as pd
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from datetime import datetime

    # Read CSV data
    df = pd.read_csv(csv_path)

    # Create a new Document
    doc = Document()

    # Set document properties
    doc.core_properties.title = "Property Inspection Report"
    doc.core_properties.author = "PhillyScript"

    # Format the inspection date
    try:
        # Parse the date from the form (YYYY-MM-DD)
        inspection_date_obj = datetime.strptime(inspection_date, '%Y-%m-%d')
        # Format as Month Day, Year
        formatted_date = inspection_date_obj.strftime('%B %d, %Y')
    except:
        # Fallback to the provided string if parsing fails
        formatted_date = inspection_date

    # REPLACE ALL THIS BOILERPLATE CODE:
    # if report_type == 'inventory':
    #     # INVENTORY CHECK-IN REPORT FORMAT
    #     title = doc.add_heading('INVENTORY', 0)
    #     ...
    # WITH THIS SINGLE LINE:
    generate_report_boilerplate(doc, report_type, address, formatted_date, on_behalf_of)

    # Group the data by room - using a simple approach to handle empty room cells
    room_data = {}
    current_room = None

    # First pass - group rows by room
    for _, row in df.iterrows():
        room = row['Room'].strip() if pd.notna(row['Room']) and row['Room'].strip() else None

        if room:
            current_room = room

        if current_room not in room_data:
            room_data[current_room] = []

        room_data[current_room].append(row)

    # Update the room processing in generate_enhanced_docx_report function
    # This restructures how we add room content to ensure images come at the end

    # Second pass - process each room group
    first_room = True
    room_counter = 1  # Counter for room numbering

    for room, rows in room_data.items():
        if room is None:
            continue  # Skip rows with no room (should be rare/nonexistent)

        # Add a page break except for the first room
        if not first_room:
            doc.add_page_break()
        first_room = False

        # Add room heading with number
        numbered_room_heading = f"{room_counter}. {room}"
        room_heading = doc.add_heading(numbered_room_heading, level=1)
        room_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Create a table for this room's inventory items
        # Columns: Attribute, Feature, Comment, Tenant Responsibility (TR)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set the header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Attribute'
        header_cells[1].text = 'Feature'
        header_cells[2].text = 'Comment'
        header_cells[3].text = 'Tenant Responsibility (TR)'

        # Apply header formatting
        for cell in header_cells:
            cell_para = cell.paragraphs[0]
            cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell_para.runs:
                run.bold = True

        # Add the data rows for this room
        # Track attributes to add numbering
        attribute_counter = 1
        current_attribute = None

        for row_data in rows:
            # Get attribute value
            attribute = row_data['Attribute'] if pd.notna(row_data['Attribute']) else ''

            # Update attribute numbering if this is a new attribute
            if attribute and attribute != current_attribute:
                current_attribute = attribute
                numbered_attribute = f"{room_counter}.{attribute_counter} {attribute}"
                attribute_counter += 1
            else:
                # Empty string if this row doesn't have an attribute (continuing from previous)
                numbered_attribute = ''

            # Add a new row to the table
            new_row = table.add_row().cells

            # Fill in the cells
            new_row[0].text = numbered_attribute
            new_row[1].text = row_data['Feature'] if pd.notna(row_data['Feature']) else ''
            new_row[2].text = row_data['Comment'] if pd.notna(row_data['Comment']) else ''
            new_row[3].text = row_data['Tenant Responsibility (TR)'] if pd.notna(
                row_data['Tenant Responsibility (TR)']) else ''

            # Center align the TR column
            if new_row[3].text:
                new_row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # NOW add room images AFTER the inventory table - this ensures they appear at the end of the room section
        # This is the key change from the previous version
        room_lower = room.lower() if isinstance(room, str) else ""
        logging.info(f"Looking for images for room: {room_lower}")
        logging.info(f"Available room keys in room_images: {list(room_images.keys())}")

        if room_lower in room_images and room_images[room_lower]:
            image_count = len(room_images[room_lower])
            logging.info(f"Found {image_count} images for room: {room_lower}")

            # Add images heading if we have images
            if image_count > 0:
                img_heading = doc.add_heading('Room Images', level=2)

            # Add each image in sequence
            for i, img_path in enumerate(room_images[room_lower]):
                logging.info(f"Adding image {i + 1}: {img_path}")
                try:
                    # Create a paragraph for the image
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Verify the image path exists
                    if not os.path.exists(img_path):
                        logging.error(f"Image file not found: {img_path}")
                        p.add_run(f"[Image {i + 1} file not found]")
                        continue

                    # Log image file details
                    file_size = os.path.getsize(img_path)
                    logging.info(f"Image file size: {file_size} bytes")

                    # Add the image to the paragraph
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(5.0))  # Standard width

                    # Add caption
                    caption = doc.add_paragraph(f"Image {i + 1}")
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.style = 'Caption'

                    # Add some space
                    doc.add_paragraph()

                    logging.info(f"Successfully added image {i + 1}")

                except Exception as e:
                    # If adding the image fails, add a placeholder text
                    error_p = doc.add_paragraph(f"[Image {i + 1} could not be displayed: {str(e)}]")
                    error_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    logging.error(f"Error adding image to report: {str(e)}", exc_info=True)

                    # Add space even if the image fails
                    doc.add_paragraph()
        else:
            logging.info(f"No images found for room: {room_lower}")

        # Increment room counter for the next room
        room_counter += 1

    # REPLACE ALL THIS CLOSING CODE:
    # # Add a page break before closing sections
    # doc.add_page_break()
    #
    # # Add closing content based on report type
    # if report_type == 'inventory':
    #     # Add tenant and landlord declaration section
    #     ...
    # WITH THIS SINGLE LINE:
    generate_report_closing(doc, report_type, 124)  # 124 is a placeholder for page count

    # Save the document with report type in the filename
    result_id = str(uuid.uuid4())
    filename = f"{result_id}_{report_type}_report.docx"
    result_path = RESULT_FOLDER / filename
    doc.save(result_path)

    # Clean up temporary image files
    for room_img_list in room_images.values():
        for img_path in room_img_list:
            try:
                Path(img_path).unlink(missing_ok=True)
            except:
                pass

    return result_path

@application.errorhandler(Exception)
def handle_error(e):
    logger.error(f"Unhandled exception: {str(e)}")
    return jsonify({
        'error': 'Internal server error',
        'message': str(e)
    }), 500

if __name__ == '__main__':
    # Determine optimal number of threads
    num_threads = min(multiprocessing.cpu_count() * 2, 8)  # Use 2x CPU cores, max 8

    # Get port from environment or use 8080 as default for App Runner
    port = int(os.environ.get('PORT', 5001))

    print(f"Starting PhillyScript server with {num_threads} worker threads on port {port}...")
    print("Available routes:")
    print(f"  - http://127.0.0.1:{port}/ (main interface)")
    print(f"  - http://127.0.0.1:{port}/transcribe (transcription page)")
    print(f"  - http://127.0.0.1:{port}/diff_check (image comparison)")

    # Set threaded mode and optimal worker count
    application.run(debug=False, host='0.0.0.0', port=port, threaded=True)